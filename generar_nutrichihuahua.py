#!/usr/bin/env python3
"""
generar_nutrichihuahua.py
Reporte NutriChihuahua — mismo formato que motor_reporte_padron.py
Uso: python3 generar_nutrichihuahua.py <excel_path> <mes> <año> <output_path>
"""
import sys, os, unicodedata, datetime
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

EXCEL    = sys.argv[1] if len(sys.argv) > 1 else ''
MES      = sys.argv[2] if len(sys.argv) > 2 else 'Enero'
ANO      = sys.argv[3] if len(sys.argv) > 3 else '2026'
OUT_PATH = sys.argv[4] if len(sys.argv) > 4 else f'NutriChihuahua_{MES}_{ANO}.docx'

RANGOS = ['0-5','6-11','12-17','18-29','30-49','50-64','65+']
RLAB   = {'0-5':'0–5','6-11':'6–11','12-17':'12–17','18-29':'18–29',
          '30-49':'30–49','50-64':'50–64','65+':'65+'}
TITULO_CORTO = f'NutriChihuahua — {MES} {ANO}'

# ─── COLORES (idénticos a motor_reporte_padron) ───────────────────────────────
AZUL_GOB  = RGBColor(0x1B, 0x3A, 0x6B)
AZUL_MED  = RGBColor(0x2E, 0x5B, 0xA8)
AZUL_CLAR = RGBColor(0xD6, 0xE4, 0xF7)
GRIS_TEXT = RGBColor(0x33, 0x33, 0x33)
BLANCO    = RGBColor(0xFF, 0xFF, 0xFF)
DORADO    = RGBColor(0xC8, 0xA0, 0x00)

# ─── HELPERS (copiados exactos de motor_reporte_padron.py) ────────────────────
def fmt(n, dec=0):
    if n is None or n == '': return '0'
    try:
        num = float(n)
        if dec == 0: return f'{int(round(num)):,}'
        return f'{num:,.{dec}f}'
    except: return '0'

def pct_of(part, total):
    if not total or total == 0: return '0'
    try: return f'{float(part)/float(total)*100:.1f}%'
    except: return '0'

def sf(v):
    try: return float(v) if v is not None else 0.0
    except: return 0.0

_MINUSCULAS = {'a','ante','bajo','con','contra','de','del','desde','durante','el','en',
               'entre','hacia','hasta','la','las','lo','los','mediante','para','por',
               'que','se','sin','sobre','su','sus','tras','un','una','unas','unos','y'}
def tc(s):
    if not s: return s
    if s != s.upper(): return s
    words = s.split()
    return ' '.join(w.capitalize() if (i==0 or w.lower() not in _MINUSCULAS) else w.lower()
                    for i,w in enumerate(words))

def rgb_hex(c): return f'{c[0]:02X}{c[1]:02X}{c[2]:02X}'

def set_cell_bg(cell, color):
    tc2 = cell._tc; tcPr = tc2.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),'clear'); shd.set(qn('w:color'),'auto')
    shd.set(qn('w:fill'), rgb_hex(color)); tcPr.append(shd)

def set_cell_borders(cell, color='CCCCCC'):
    tc2 = cell._tc; tcPr = tc2.get_or_add_tcPr()
    tcB = OxmlElement('w:tcBorders')
    for side in ['top','left','bottom','right']:
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'),'single'); el.set(qn('w:sz'),'4')
        el.set(qn('w:space'),'0');    el.set(qn('w:color'),color)
        tcB.append(el)
    tcPr.append(tcB)

def set_cell_no_borders(cell):
    tc2 = cell._tc; tcPr = tc2.get_or_add_tcPr()
    tcB = OxmlElement('w:tcBorders')
    for side in ['top','left','bottom','right']:
        el = OxmlElement(f'w:{side}'); el.set(qn('w:val'),'none'); tcB.append(el)
    tcPr.append(tcB)

def set_cell_margins(cell, top=80, bottom=80, left=120, right=120):
    tc2 = cell._tc; tcPr = tc2.get_or_add_tcPr()
    tcM = OxmlElement('w:tcMar')
    for side, val in [('top',top),('bottom',bottom),('left',left),('right',right)]:
        el = OxmlElement(f'w:{side}'); el.set(qn('w:w'),str(val)); el.set(qn('w:type'),'dxa')
        tcM.append(el)
    tcPr.append(tcM)

def set_col_width(cell, w):
    tc2 = cell._tc; tcPr = tc2.get_or_add_tcPr()
    tcW = OxmlElement('w:tcW'); tcW.set(qn('w:w'),str(w)); tcW.set(qn('w:type'),'dxa')
    tcPr.append(tcW)

def prevent_row_break(row):
    tr = row._tr; trPr = tr.get_or_add_trPr()
    cs = OxmlElement('w:cantSplit'); cs.set(qn('w:val'),'1'); trPr.append(cs)

def keep_with_next_row(row):
    for cell in row.cells:
        for p in cell.paragraphs:
            pPr = p._p.get_or_add_pPr()
            kn = OxmlElement('w:keepNext'); kn.set(qn('w:val'),'1'); pPr.append(kn)

def set_page_break_before(p):
    pPr = p._p.get_or_add_pPr()
    for existing in pPr.findall(qn('w:pageBreakBefore')): pPr.remove(existing)
    pbr = OxmlElement('w:pageBreakBefore'); pbr.set(qn('w:val'),'1'); pPr.insert(0, pbr)

def add_bottom_border(p, color='2E5BA8', sz='6'):
    pPr = p._p.get_or_add_pPr(); pBdr = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'),'single'); bot.set(qn('w:sz'),sz)
    bot.set(qn('w:space'),'4'); bot.set(qn('w:color'),color)
    pBdr.append(bot); pPr.append(pBdr)

def add_top_border(p, color='1B3A6B', sz='4'):
    pPr = p._p.get_or_add_pPr(); pBdr = OxmlElement('w:pBdr')
    top_el = OxmlElement('w:top')
    top_el.set(qn('w:val'),'single'); top_el.set(qn('w:sz'),sz)
    top_el.set(qn('w:space'),'4'); top_el.set(qn('w:color'),color)
    pBdr.append(top_el); pPr.append(pBdr)

def add_heading(doc, text, page_break=True):
    p = doc.add_paragraph()
    if page_break: set_page_break_before(p)
    run = p.add_run(text)
    run.font.name='Arial'; run.font.size=Pt(16); run.font.bold=True
    run.font.color.rgb=AZUL_GOB
    p.paragraph_format.space_before=Pt(0); p.paragraph_format.space_after=Pt(4)
    add_bottom_border(p)

def add_subheading(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name='Arial'; run.font.size=Pt(11); run.font.bold=True
    run.font.color.rgb=AZUL_MED
    p.paragraph_format.space_before=Pt(8); p.paragraph_format.space_after=Pt(3)

def add_body(doc, text, size=10, color=None, italic=False, keep_next=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name='Arial'; run.font.size=Pt(size); run.font.italic=italic
    run.font.color.rgb = color if color else GRIS_TEXT
    p.paragraph_format.space_before=Pt(2); p.paragraph_format.space_after=Pt(2)
    if keep_next:
        pPr = p._p.get_or_add_pPr()
        kn = OxmlElement('w:keepNext'); kn.set(qn('w:val'),'1'); pPr.append(kn)

def add_spacer(doc, pts=6, keep_next=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before=Pt(pts); p.paragraph_format.space_after=Pt(0)
    if keep_next:
        pPr = p._p.get_or_add_pPr()
        kn = OxmlElement('w:keepNext'); kn.set(qn('w:val'),'1'); pPr.append(kn)

def set_repeat_header(row):
    tr = row._tr; trPr = tr.get_or_add_trPr()
    tblH = OxmlElement('w:tblHeader'); tblH.set(qn('w:val'),'1'); trPr.append(tblH)

def set_table_width(table, width_twips):
    tbl = table._tbl; tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None: tblPr = OxmlElement('w:tblPr'); tbl.insert(0, tblPr)
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), str(width_twips)); tblW.set(qn('w:type'),'dxa')
    existing = tblPr.find(qn('w:tblW'))
    if existing is not None: tblPr.remove(existing)
    tblPr.append(tblW)

def add_table(doc, headers, rows, col_widths, first_col_bold_size=None):
    PAGE_W = 10640
    raw_sum = sum(col_widths)
    scaled = [int(w * PAGE_W / raw_sum) for w in col_widths]
    scaled[-1] += PAGE_W - sum(scaled)
    col_widths = scaled
    n = len(headers)
    t = doc.add_table(rows=1+len(rows), cols=n)
    t.style='Table Grid'; t.alignment=WD_TABLE_ALIGNMENT.LEFT
    set_table_width(t, PAGE_W)
    hr = t.rows[0]; set_repeat_header(hr)
    for ci, (h, w) in enumerate(zip(headers, col_widths)):
        cell = hr.cells[ci]
        set_cell_bg(cell, AZUL_GOB); set_cell_borders(cell,'1B3A6B')
        set_cell_margins(cell); set_col_width(cell, w)
        p = cell.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(str(h) if h else '')
        run.font.name='Arial'; run.font.size=Pt(7.5); run.font.bold=True
        run.font.color.rgb=BLANCO
        p.paragraph_format.space_before=Pt(0); p.paragraph_format.space_after=Pt(0)
        cell.vertical_alignment=WD_ALIGN_VERTICAL.CENTER
    n_rows = len(rows)
    for ri, row_data in enumerate(rows):
        is_total = str(row_data[0]).strip().upper().startswith('TOTAL')
        is_last  = (ri == n_rows - 1)
        bg = AZUL_CLAR if is_total else (BLANCO if ri%2==0 else AZUL_CLAR)
        row = t.rows[ri+1]; prevent_row_break(row)
        if not is_last and ri == n_rows - 2: keep_with_next_row(row)
        for ci, (val, w) in enumerate(zip(row_data, col_widths)):
            cell = row.cells[ci]
            set_cell_bg(cell, bg); set_cell_borders(cell)
            set_cell_margins(cell); set_col_width(cell, w)
            p = cell.paragraphs[0]
            p.alignment=WD_ALIGN_PARAGRAPH.LEFT if ci==0 else WD_ALIGN_PARAGRAPH.RIGHT
            run = p.add_run('0' if val is None else str(val))
            run.font.name='Arial'
            if not is_total and ci==0 and first_col_bold_size:
                run.font.size=Pt(first_col_bold_size); run.font.bold=True
            else:
                run.font.size=Pt(9) if is_total else Pt(7.5)
                run.font.bold=is_total
            run.font.color.rgb=AZUL_GOB if is_total else GRIS_TEXT
            p.paragraph_format.space_before=Pt(0); p.paragraph_format.space_after=Pt(0)
            cell.vertical_alignment=WD_ALIGN_VERTICAL.CENTER

def add_kpi_table(doc, kpi_rows):
    for group in kpi_rows:
        KPI_W = 10640
        n_cols = len(group)
        if n_cols == 0: continue
        t = doc.add_table(rows=1, cols=n_cols); t.alignment=WD_TABLE_ALIGNMENT.LEFT
        set_table_width(t, KPI_W)
        col_w = KPI_W // n_cols
        row = t.rows[0]
        for ci, kpi in enumerate(group):
            cell = row.cells[ci]
            set_cell_bg(cell, AZUL_CLAR); set_cell_no_borders(cell)
            set_cell_margins(cell,80,80,120,120); set_col_width(cell, col_w)
            p1 = cell.paragraphs[0]; p1.alignment=WD_ALIGN_PARAGRAPH.CENTER
            r1 = p1.add_run(kpi.get('value','-'))
            r1.font.name='Arial'; r1.font.size=Pt(15); r1.font.bold=True
            r1.font.color.rgb=AZUL_GOB
            p1.paragraph_format.space_before=Pt(0); p1.paragraph_format.space_after=Pt(2)
            p2 = cell.add_paragraph(); p2.alignment=WD_ALIGN_PARAGRAPH.CENTER
            r2 = p2.add_run(kpi.get('label',''))
            r2.font.name='Arial'; r2.font.size=Pt(7.5); r2.font.bold=True
            r2.font.color.rgb=GRIS_TEXT
            p2.paragraph_format.space_before=Pt(0); p2.paragraph_format.space_after=Pt(2)
            if kpi.get('sub'):
                p3 = cell.add_paragraph(); p3.alignment=WD_ALIGN_PARAGRAPH.CENTER
                r3 = p3.add_run(kpi['sub'])
                r3.font.name='Arial'; r3.font.size=Pt(6.5); r3.font.italic=True
                r3.font.color.rgb=RGBColor(0x66,0x66,0x66)
                p3.paragraph_format.space_before=Pt(0); p3.paragraph_format.space_after=Pt(0)
        add_spacer(doc, 4)

def set_page_margins(section):
    section.top_margin=Cm(1.8); section.bottom_margin=Cm(1.8)
    section.left_margin=Cm(1.41); section.right_margin=Cm(1.41)

def add_header_footer(section):
    hdr = section.header
    hp  = hdr.paragraphs[0] if hdr.paragraphs else hdr.add_paragraph()
    hp.clear()
    run = hp.add_run(f'Reporte — {TITULO_CORTO}  |  Gobierno del Estado de Chihuahua')
    run.font.name='Arial'; run.font.size=Pt(7.5); run.font.color.rgb=AZUL_GOB
    add_bottom_border(hp, color='1B3A6B', sz='4')
    ftr = section.footer
    fp  = ftr.paragraphs[0] if ftr.paragraphs else ftr.add_paragraph()
    fp.clear(); fp.alignment=WD_ALIGN_PARAGRAPH.LEFT
    run = fp.add_run('Secretaría de Desarrollo Humano y Bien Común  —  Chihuahua, México')
    run.font.name='Arial'; run.font.size=Pt(7); run.font.color.rgb=RGBColor(0x88,0x88,0x88)
    add_top_border(fp)

def build_portada(doc, fecha_str):
    section = doc.sections[0]
    section.top_margin=Cm(2.5); section.bottom_margin=Cm(2.5)
    section.left_margin=Cm(2.0); section.right_margin=Cm(2.0)
    section.different_first_page_header_footer=True

    for _ in range(6): add_spacer(doc, 8)

    p = doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('GOBIERNO DEL ESTADO DE CHIHUAHUA')
    r.font.name='Arial'; r.font.size=Pt(13); r.font.bold=True; r.font.color.rgb=AZUL_GOB
    p.paragraph_format.space_after=Pt(4)

    p = doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('Secretaría de Desarrollo Humano y Bien Común')
    r.font.name='Arial'; r.font.size=Pt(11); r.font.color.rgb=AZUL_MED
    p.paragraph_format.space_after=Pt(18)

    p = doc.add_paragraph(); add_bottom_border(p, color='1B3A6B', sz='12')
    p.paragraph_format.space_after=Pt(22)

    p = doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('REPORTE DE AVANCE')
    r.font.name='Arial'; r.font.size=Pt(22); r.font.bold=True; r.font.color.rgb=AZUL_GOB
    p.paragraph_format.space_after=Pt(8)

    p = doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('Programa NutriChihuahua')
    r.font.name='Arial'; r.font.size=Pt(14); r.font.color.rgb=AZUL_MED
    p.paragraph_format.space_after=Pt(16)

    p = doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f'{MES.upper()} {ANO}')
    r.font.name='Arial'; r.font.size=Pt(28); r.font.bold=True; r.font.color.rgb=DORADO
    p.paragraph_format.space_after=Pt(22)

    p = doc.add_paragraph(); add_bottom_border(p, color='1B3A6B', sz='12')
    p.paragraph_format.space_after=Pt(24)

    for _ in range(2): add_spacer(doc, 10)

    p = doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f'Período reportado: {MES} {ANO}')
    r.font.name='Arial'; r.font.size=Pt(11); r.font.bold=True; r.font.color.rgb=AZUL_MED
    p.paragraph_format.space_after=Pt(6)

    p = doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f'Fecha de emisión: {fecha_str}')
    r.font.name='Arial'; r.font.size=Pt(10); r.font.color.rgb=GRIS_TEXT

# ─── LECTURA DE DATOS ─────────────────────────────────────────────────────────
def leer_nutri(excel_path):
    import pandas as pd
    import warnings; warnings.filterwarnings('ignore')
    xl = pd.ExcelFile(str(excel_path))
    if 'Nutrichihuahua' not in xl.sheet_names:
        raise ValueError('No existe la hoja Nutrichihuahua en el Excel.')
    df = xl.parse('Nutrichihuahua', header=None)

    # Columnas descubiertas dinámicamente: un refresh de Excel puede reordenar o
    # insertar sub-bloques (p.ej. añadió un bloque "Sin datos" que no existía
    # antes) y romper cualquier offset fijo tipo "col 2+j".
    _m_col_map, _h_col_map = {}, {}
    _col_total_m = _col_total_h = None
    for _sri in range(min(15, df.shape[0]) - 1):
        _sect = [str(v).strip() if isinstance(v, str) else '' for v in df.iloc[_sri]]
        if 'M' in _sect and 'H' in _sect and 'Total M' in _sect and 'Total H' in _sect:
            _rango_row = [str(v).strip() if isinstance(v, str) else '' for v in df.iloc[_sri + 1]]
            _m0, _m1 = _sect.index('M'), _sect.index('Total M')
            _h0, _h1 = _sect.index('H'), _sect.index('Total H')
            _m_col_map = {_rango_row[ci]: ci for ci in range(_m0, _m1) if _rango_row[ci] in RANGOS}
            _h_col_map = {_rango_row[ci]: ci for ci in range(_h0, _h1) if _rango_row[ci] in RANGOS}
            _col_total_m, _col_total_h = _m1, _h1
            break
    if not _m_col_map or not _h_col_map:
        _m_col_map = {r: 2 + j for j, r in enumerate(RANGOS)}
        _h_col_map = {r: 11 + j for j, r in enumerate(RANGOS)}
        _col_total_m, _col_total_h = 10, 19

    muns = []
    for row_i in range(8, df.shape[0]):
        mun = str(df.iloc[row_i, 1]).strip() if not __import__('pandas').isna(df.iloc[row_i, 1]) else ''
        if not mun or mun.upper() in ('NAN','NONE','(EN BLANCO)','GRAND TOTAL','TOTAL','FORANEO'): continue
        def _iv(col):
            v = df.iloc[row_i, col]
            return int(v) if not __import__('pandas').isna(v) else 0
        total_m  = _iv(_col_total_m)
        m_rangos = {r: _iv(ci) for r, ci in _m_col_map.items()}
        h_rangos = {r: _iv(ci) for r, ci in _h_col_map.items()}
        total_h  = _iv(_col_total_h)
        muns.append({'nombre': mun, 'total': total_m+total_h,
                     'mujeres': total_m, 'hombres': total_h,
                     'rangos_m': m_rangos, 'rangos_h': h_rangos})

    # Columnas de la tabla AH (institución→programa→apoyo) descubiertas
    # dinámicamente: mismo problema de sub-bloques reordenados/insertados
    # que en la Tabla 1 y la tabla CC.
    _ah_m_map, _ah_h_map = {}, {}
    _ah_col_tm = _ah_col_th = _ah_col_tot = None
    for _sri in range(min(15, df.shape[0]) - 1):
        _region = {ci: (str(v).strip() if isinstance(v, str) else '')
                   for ci, v in enumerate(df.iloc[_sri]) if 30 <= ci < 65}
        _vals = set(_region.values())
        if {'M', 'H', 'Total M', 'Total H', 'TOTAL'} <= _vals:
            _rango_row = {ci: (str(v).strip() if isinstance(v, str) else '')
                          for ci, v in enumerate(df.iloc[_sri + 1]) if 30 <= ci < 65}
            _m0 = next(ci for ci, v in _region.items() if v == 'M')
            _ah_col_tm = next(ci for ci, v in _region.items() if v == 'Total M')
            _h0 = next(ci for ci, v in _region.items() if v == 'H')
            _ah_col_th = next(ci for ci, v in _region.items() if v == 'Total H')
            _ah_col_tot = next(ci for ci, v in _region.items() if v == 'TOTAL')
            _ah_m_map = {v: ci for ci, v in _rango_row.items() if _m0 <= ci < _ah_col_tm and v in RANGOS}
            _ah_h_map = {v: ci for ci, v in _rango_row.items() if _h0 <= ci < _ah_col_th and v in RANGOS}
            break
    if _ah_col_tot is None:
        _ah_m_map = {r: 34 + j for j, r in enumerate(RANGOS)}
        _ah_h_map = {r: 43 + j for j, r in enumerate(RANGOS)}
        _ah_col_tm, _ah_col_th, _ah_col_tot = 42, 51, 52

    INST_NAMES = ['DIF','SDHyBC','SPyCI']
    insts, cur = [], None
    total_benef_global = total_m_global = total_h_global = 0
    for row_i in range(9, df.shape[0]):
        prog = str(df.iloc[row_i, 33]).strip() if not __import__('pandas').isna(df.iloc[row_i, 33]) else ''
        if not prog or prog in ['nan','None','Grand Total','LOCALIZABLES NUTRICHIHUAHUA']: continue
        tv = df.iloc[row_i, _ah_col_tot]
        if __import__('pandas').isna(tv): continue
        total = int(tv)
        if prog.upper() == 'TOTAL':
            total_benef_global = total
            total_m_global = int(df.iloc[row_i, _ah_col_tm]) if not __import__('pandas').isna(df.iloc[row_i, _ah_col_tm]) else 0
            total_h_global = int(df.iloc[row_i, _ah_col_th]) if not __import__('pandas').isna(df.iloc[row_i, _ah_col_th]) else 0
            continue
        if total == 0: continue
        def _iv2(col):
            v = df.iloc[row_i, col]
            return int(v) if not __import__('pandas').isna(v) else 0
        entry = {'nombre': prog, 'total': total,
                 'mujeres': _iv2(_ah_col_tm), 'hombres': _iv2(_ah_col_th),
                 'rangos_m': {r: _iv2(ci) for r, ci in _ah_m_map.items()},
                 'rangos_h': {r: _iv2(ci) for r, ci in _ah_h_map.items()},
                 'programas': []}
        if prog in INST_NAMES:
            cur = entry; insts.append(entry)
        elif prog != 'TOTAL' and cur:
            cur['programas'].append(entry)

    # ── FIX: la tabla AH también tiene jerarquía de 3 niveles ──────────────────
    # institución → PROGRAMA → apoyo (igual que la tabla BE más abajo).
    # El parseo de arriba mete TODAS las filas no-institución directo en
    # inst['programas'] como si fueran programas hermanos, pero algunas son en
    # realidad apoyos (sub-filas de un programa). Esto causaba que un mismo
    # nombre de apoyo (p.ej. "Asistencia Alimentaria en Espacio Comun") apareciera
    # una vez por cada programa padre al que pertenece, en vez de filtrarse.
    # Se reutiliza la misma heurística de totales que la tabla BE: una fila es
    # PROGRAMA si su total ≈ la suma acumulada de las filas que le siguen
    # (antes de que la acumulación se dispare muy por encima de su total).
    def _es_programa_ah(idx, rows):
        total = rows[idx]['total']
        acum = 0
        for r in rows[idx+1:]:
            acum += r['total']
            if abs(acum - total) <= max(5, int(total * 0.02)):
                return True
            if acum > total * 1.05:
                break
        return False

    for inst_entry in insts:
        raw_progs = inst_entry['programas']
        inst_entry['programas'] = [
            p for idx, p in enumerate(raw_progs) if _es_programa_ah(idx, raw_progs)
        ]

    # ── Tablas de apoyos entregados ──────────────────────────────────────────
    # Excel mueve estos pivotes cuando se agregan bloques como "Sin datos".
    # Se detectan por sus anclas, no por columnas fijas:
    #   1) institución → programa → tipo de apoyo
    #   2) municipio → total de apoyos
    def _count_blocks():
        blocks = []
        for anchor_row in range(df.shape[0]):
            for label_col, value in enumerate(df.iloc[anchor_row]):
                if str(value).strip().upper() != 'RECUENTO DE # DE APOYOS':
                    continue
                data_label_row = None
                for ri in range(anchor_row + 1, min(anchor_row + 5, df.shape[0])):
                    if str(df.iloc[ri, label_col]).strip().upper() == 'LOCALIZABLES NUTRICHIHUAHUA':
                        data_label_row = ri
                        break
                if data_label_row is None:
                    continue

                metric_cols = None
                for ri in range(anchor_row, data_label_row + 1):
                    region = {ci: str(df.iloc[ri, ci]).strip()
                              for ci in range(label_col, min(label_col + 32, df.shape[1]))
                              if isinstance(df.iloc[ri, ci], str)}
                    vals = set(region.values())
                    if {'M', 'H', 'Total M', 'Total H', 'TOTAL'} <= vals:
                        metric_cols = (
                            next(ci for ci, v in region.items() if v == 'Total M'),
                            next(ci for ci, v in region.items() if v == 'Total H'),
                            next(ci for ci, v in region.items() if v == 'TOTAL'),
                        )
                        break
                if metric_cols:
                    blocks.append((label_col, data_label_row + 1, *metric_cols))
        return blocks

    count_blocks = _count_blocks()

    # Jerarquía institución → programa → apoyo. Se reconoce porque la primera
    # fila de datos es una institución válida.
    insts_ap = []
    for label_col, data_start, col_tm, col_th, col_tot in count_blocks:
        first_name = str(df.iloc[data_start, label_col]).strip() if data_start < df.shape[0] else ''
        if first_name not in INST_NAMES:
            continue
        cur_ap = None
        for row_i in range(data_start, df.shape[0]):
            name = str(df.iloc[row_i, label_col]).strip() if pd.notna(df.iloc[row_i, label_col]) else ''
            if not name or name in ('nan', 'None', 'Grand Total'):
                continue
            if name.upper() == 'TOTAL':
                break
            tv = df.iloc[row_i, col_tot]
            if pd.isna(tv) or int(tv) == 0:
                continue
            entry = {
                'nombre': name,
                'total': int(tv),
                'mujeres': int(df.iloc[row_i, col_tm]) if pd.notna(df.iloc[row_i, col_tm]) else 0,
                'hombres': int(df.iloc[row_i, col_th]) if pd.notna(df.iloc[row_i, col_th]) else 0,
            }
            if name in INST_NAMES:
                entry['programas_raw'] = []
                cur_ap = entry
                insts_ap.append(entry)
            elif cur_ap:
                cur_ap['programas_raw'].append(entry)

        for inst_entry in insts_ap:
            rows = inst_entry.pop('programas_raw', [])
            prog_idxs = [idx for idx in range(len(rows)) if _es_programa_ah(idx, rows)]
            programas = []
            for pos, idx in enumerate(prog_idxs):
                next_idx = prog_idxs[pos + 1] if pos + 1 < len(prog_idxs) else len(rows)
                programas.append({**rows[idx], 'apoyos': rows[idx + 1:next_idx]})
            inst_entry['programas'] = programas
        break

    # Respaldo para layouts antiguos que no tengan el pivote de apoyos.
    if not insts_ap:
        insts_ap = [
            {**inst_entry, 'programas': [
                {**p, 'apoyos': []} for p in inst_entry['programas']
            ]}
            for inst_entry in insts
        ]

    # Tabla de apoyos por municipio: es el otro bloque cuyo primer dato no es
    # una institución. Se usa para el reporte y sus KPIs municipales.
    SKIP_MUN = {'nan','None','Grand Total','FORANEO','NO IDENTIFICADO','TOTAL'}
    muns_ap = []
    for label_col, data_start, col_tm, col_th, col_tot in count_blocks:
        first_name = str(df.iloc[data_start, label_col]).strip() if data_start < df.shape[0] else ''
        if first_name in INST_NAMES:
            continue
        for row_i in range(data_start, df.shape[0]):
            mun = str(df.iloc[row_i, label_col]).strip() if pd.notna(df.iloc[row_i, label_col]) else ''
            if not mun or mun in SKIP_MUN:
                if mun == 'TOTAL':
                    break
                continue
            tv = df.iloc[row_i, col_tot]
            if pd.isna(tv) or int(tv) == 0:
                continue
            muns_ap.append({
                'nombre': mun,
                'total': int(tv),
                'mujeres': int(df.iloc[row_i, col_tm]) if pd.notna(df.iloc[row_i, col_tm]) else 0,
                'hombres': int(df.iloc[row_i, col_th]) if pd.notna(df.iloc[row_i, col_th]) else 0,
            })
        if muns_ap:
            break
    # La fila TOTAL deduplica a las personas que aparecen en más de una
    # institución. Sumar instituciones inflaba el universo localizable.
    total_benef_canonico = total_benef_global
    if not total_benef_canonico:
        total_benef_canonico = sum(i['total'] for i in insts)
    apoyos_total = sum(i['total'] for i in insts_ap)
    if not apoyos_total:
        # Igual que arriba: la tabla BE (apoyos por institución) también
        # desapareció; se usa la Tabla CC (apoyos por municipio) como respaldo.
        apoyos_total = sum(m['total'] for m in muns_ap)

    return {
        'municipios':    [m for m in muns if m['nombre'] not in ('TOTAL','NO IDENTIFICADO')],
        'instituciones': insts,
        'apoyos_inst':   insts_ap,
        'apoyos_mun':    muns_ap,
        'apoyos_total':  apoyos_total,
        'total_benef_canonico': total_benef_canonico,
        'total_m_canonico': total_m_global,
        'total_h_canonico': total_h_global,
    }

# ─── MAIN ─────────────────────────────────────────────────────────────────────
def main():
    if not EXCEL or not Path(EXCEL).exists():
        print(f'ERROR: No se encontró el archivo Excel: {EXCEL}', file=sys.stderr); sys.exit(1)

    # ── Modo --dashboard-only: solo actualiza js_render_nutri.js ─────────────
    if '--dashboard-only' in sys.argv:
        import json, re as _re
        print('Modo --dashboard-only: regenerando js_render_nutri.js...')
        try:
            data = leer_nutri(EXCEL)
        except Exception as e:
            print(f'ERROR: {e}', file=sys.stderr); sys.exit(1)
        insts_ap  = data['apoyos_inst']
        insts_real = data['instituciones']  # tabla AH — fuente correcta de BENEFICIARIOS (no apoyos)
        # Lookup de beneficiarios reales por nombre de institución
        _benef_by_inst = {i['nombre']: i for i in insts_real}
        _ap_by_inst = {i['nombre']: i for i in insts_ap}
        muns      = data['municipios']
        RLAB_D    = {'0-5':'0–5','6-11':'6–11','12-17':'12–17','18-29':'18–29',
                     '30-49':'30–49','50-64':'50–64','65+':'65+'}
        RT_M = {r: sum(m['rangos_m'].get(r,0) for m in muns) for r in RANGOS}
        RT_H = {r: sum(m['rangos_h'].get(r,0) for m in muns) for r in RANGOS}
        RT   = {r: RT_M[r]+RT_H[r] for r in RANGOS}
        # Apoyos desde insts_ap
        apoyos_dict = {}
        for inst in insts_ap:
            for prog in inst.get('programas',[]):
                for ap in prog.get('apoyos',[]):
                    k = ap['nombre']
                    if k not in apoyos_dict:
                        apoyos_dict[k] = {'n':k,'t':0,'m':0,'h':0,'insts':[]}
                    apoyos_dict[k]['t'] += ap['total']
                    apoyos_dict[k]['m'] += ap['mujeres']
                    apoyos_dict[k]['h'] += ap['hombres']
                    if inst['nombre'] not in apoyos_dict[k]['insts']:
                        apoyos_dict[k]['insts'].append(inst['nombre'])
        apoyos_list = sorted(apoyos_dict.values(), key=lambda x:-x['t'])
        total_benef  = data.get('total_benef_canonico', sum(i['total'] for i in data['instituciones']))
        total_apoyos = data['apoyos_total']
        ND = {
            'total_benef':  total_benef,
            'total_m':      data.get('total_m_canonico', 0),
            'total_h':      data.get('total_h_canonico', 0),
            'total_apoyos': total_apoyos,
            'RT':  RT, 'RANGOS': RANGOS, 'RLAB': RLAB_D,
            'muns': [{'n':m['nombre'],'t':m['total'],'m':m['mujeres'],'h':m['hombres'],
                      'at':m['total'],'am':m['mujeres'],'ah':m['hombres'],
                      'rm':m['rangos_m'],'rh':m['rangos_h']} for m in muns],
            'insts': [{
                       'nombre':i['nombre'],
                       'benef':i['total'],
                       'bm':i['mujeres'],
                       'bh':i['hombres'],
                       'apoyos_total':_ap_by_inst.get(i['nombre'], {}).get('total', 0),
                       'am':_ap_by_inst.get(i['nombre'], {}).get('mujeres', 0),
                       'ah':_ap_by_inst.get(i['nombre'], {}).get('hombres', 0),
                       'programas':[{'n':p['nombre'],'t':p['total'],'m':p['mujeres'],'h':p['hombres']}
                                    for p in i.get('programas',[])],
                       'ap_programas':[{
                           'n':p['nombre'],'t':p['total'],
                           'apoyos':[{'n':a['nombre'],'t':a['total'],'m':a['mujeres'],'h':a['hombres']}
                                     for a in p.get('apoyos',[])]
                       } for p in _ap_by_inst.get(i['nombre'], {}).get('programas',[])]
                       }
                      for i in insts_real],            'apoyos': apoyos_list,
            'RT_M': RT_M, 'RT_H': RT_H,
        }
        nd_json = json.dumps(ND, ensure_ascii=False, separators=(',',':'))
        # Actualizar js_render_nutri.js
        js_path = Path(__file__).parent / 'js_render_nutri.js'
        if not js_path.exists():
            print(f'ERROR: {js_path} no encontrado', file=sys.stderr); sys.exit(1)
        js = js_path.read_text(encoding='utf-8')
        # Reemplazar ND usando balance de llaves
        _start = js.find('const ND')
        if _start < 0:
            print('AVISO: No se encontro const ND en js_render_nutri.js', file=sys.stderr)
        else:
            _eq = js.find('{', _start)
            _depth = 0; _end = _eq
            for _i, _ch in enumerate(js[_eq:], _eq):
                if _ch == '{': _depth += 1
                elif _ch == '}': _depth -= 1
                if _depth == 0: _end = _i + 1; break
            if _end < len(js) and js[_end] == ';': _end += 1
            js_new = js[:_start] + 'const ND  = ' + nd_json + ';' + js[_end:]
            js_path.write_text(js_new, encoding='utf-8')
            print(f'  js_render_nutri.js actualizado')
            print(f'  total_benef={total_benef}, total_apoyos={total_apoyos}')
            print(f'  apoyos: {[a["n"] for a in apoyos_list]}')
        return

    print('Leyendo hoja NutriChihuahua...')
    try:
        data = leer_nutri(EXCEL)
    except Exception as e:
        print(f'ERROR: {e}', file=sys.stderr); sys.exit(1)

    muns        = data['municipios']
    insts       = data['instituciones']
    insts_ap    = data['apoyos_inst']
    muns_ap     = data['apoyos_mun']
    apoyos_total = data['apoyos_total']
    total_m_ap  = sum(i['mujeres'] for i in insts_ap)
    total_h_ap  = sum(i['hombres'] for i in insts_ap)
    total   = sum(m['total']   for m in muns)
    total_m = sum(m['mujeres'] for m in muns)
    total_h = sum(m['hombres'] for m in muns)
    mun_activos = sum(1 for m in muns if m['total'] > 0)

    RT_M = {r: sum(m['rangos_m'].get(r,0) for m in muns) for r in RANGOS}
    RT_H = {r: sum(m['rangos_h'].get(r,0) for m in muns) for r in RANGOS}
    RT   = {r: RT_M[r]+RT_H[r] for r in RANGOS}
    dom_rango = max(RANGOS, key=lambda r: RT[r])

    # ── Municipios sin cobertura vs los 67 oficiales ──────────────────────────
    MUNICIPIOS_67 = [
        'AHUMADA','ALDAMA','ALLENDE','AQUILES SERDAN','ASCENSION','BACHINIVA','BALLEZA',
        'BATOPILAS DE MANUEL GOMEZ MORIN','BOCOYNA','BUENAVENTURA','CAMARGO','CARICHI',
        'CASAS GRANDES','CHIHUAHUA','CHINIPAS','CORONADO','COYAME DEL SOTOL','CUAUHTEMOC',
        'CUSIHUIRIACHI','DELICIAS','DR. BELISARIO DOMINGUEZ','EL TULE','GALEANA',
        'GOMEZ FARIAS','GRAN MORELOS','GUACHOCHI','GUADALUPE','GUADALUPE Y CALVO',
        'GUAZAPARES','GUERRERO','HIDALGO DEL PARRAL','HUEJOTITAN','IGNACIO ZARAGOZA',
        'JANOS','JIMENEZ','JUAREZ','JULIMES','LA CRUZ','LOPEZ','MADERA','MAGUARICHI',
        'MANUEL BENAVIDES','MATACHI','MATAMOROS','MEOQUI','MORELOS','MORIS','NAMIQUIPA',
        'NONOAVA','NUEVO CASAS GRANDES','OCAMPO','OJINAGA','PRAXEDIS G. GUERRERO',
        'RIVA PALACIO','ROSALES','ROSARIO','SAN FRANCISCO DE BORJA','SAN FRANCISCO DE CONCHOS',
        'SAN FRANCISCO DEL ORO','SANTA BARBARA','SANTA ISABEL','SATEVO','SAUCILLO',
        'TEMOSACHIC','URIQUE','URUACHI','VALLE DE ZARAGOZA',
    ]
    import unicodedata as _ud
    def _norm_simple(s): return _ud.normalize('NFD',str(s).upper().strip()).encode('ascii','ignore').decode()
    muns_presentes = {_norm_simple(m['nombre']) for m in muns}
    muns_sin_cobertura = [m for m in MUNICIPIOS_67 if _norm_simple(m) not in muns_presentes]

    print(f'  Municipios: {len(muns)} | Total: {total:,} | M: {total_m:,} | H: {total_h:,}')

    meses_es = ['enero','febrero','marzo','abril','mayo','junio','julio','agosto',
                'septiembre','octubre','noviembre','diciembre']
    hoy = datetime.datetime.now()
    fecha_str = f'{hoy.day} de {meses_es[hoy.month-1]} de {hoy.year}'

    print('Construyendo documento...')
    doc = Document()
    section = doc.sections[0]
    set_page_margins(section)
    add_header_footer(section)

    # ══ PORTADA ═══════════════════════════════════════════════════════════════
    build_portada(doc, fecha_str)

    # ══ 1. RESUMEN EJECUTIVO ══════════════════════════════════════════════════
    add_heading(doc, '1. Resumen Ejecutivo', page_break=True)
    add_body(doc, (
        f'Informe de beneficiarios localizables del Programa NutriChihuahua al corte de {MES} {ANO}, '
        f'operado por la Secretaría de Desarrollo Humano y Bien Común del Gobierno del Estado de Chihuahua. '
        f'NutriChihuahua es un programa de apoyo alimentario que atiende a familias chihuahuenses '
        f'en situación de vulnerabilidad a través de DIF, SDHyBC y SPyCI.'
    ))
    add_spacer(doc, 6)

    # Calcular indicadores para KPIs
    rango_max = max(RANGOS, key=lambda r: RT[r])
    rango_min = min((r for r in RANGOS if RT[r] > 0), key=lambda r: RT[r])
    mun_ap_top = max(muns_ap, key=lambda x: x['total']) if muns_ap else None
    mun_ap_bot = min((m for m in muns_ap if m['total'] > 0), key=lambda x: x['total']) if muns_ap else None
    # Cobertura municipal = beneficiarios / población municipal
    POB_MUN = {
        'AHUMADA':16198,'ALDAMA':27591,'ALLENDE':8403,'AQUILES SERDAN':33187,
        'ASCENSION':27978,'BACHINIVA':5850,'BALLEZA':16406,
        'BATOPILAS DE MANUEL GOMEZ MORIN':11069,'BOCOYNA':23060,'BUENAVENTURA':27426,
        'CAMARGO':48426,'CARICHI':7969,'CASAS GRANDES':12513,'CHIHUAHUA':1028306,
        'CHINIPAS':5960,'CORONADO':2060,'COYAME DEL SOTOL':1218,'CUAUHTEMOC':196633,
        'CUSIHUIRIACHI':5826,'DELICIAS':156678,'DR. BELISARIO DOMINGUEZ':2475,
        'EL TULE':1369,'GALEANA':7291,'GOMEZ FARIAS':6778,'GRAN MORELOS':2484,
        'GUACHOCHI':56871,'GUADALUPE':3708,'GUADALUPE Y CALVO':50243,
        'GUAZAPARES':9305,'GUERRERO':34977,'HIDALGO DEL PARRAL':127636,
        'HUEJOTITAN':787,'IGNACIO ZARAGOZA':5040,'JANOS':11321,'JIMENEZ':39683,
        'JUAREZ':1661295,'JULIMES':5734,'LA CRUZ':3686,'LOPEZ':4291,
        'MADERA':24000,'MAGUARICHI':1277,'MANUEL BENAVIDES':1103,'MATACHI':2700,
        'MATAMOROS':4351,'MEOQUI':46611,'MORELOS':7331,'MORIS':4478,
        'NAMIQUIPA':22649,'NONOAVA':3036,'NUEVO CASAS GRANDES':68506,'OCAMPO':8965,
        'OJINAGA':24243,'PRAXEDIS G. GUERRERO':4842,'RIVA PALACIO':7722,
        'ROSALES':17031,'ROSARIO':2196,'SAN FRANCISCO DE BORJA':2315,
        'SAN FRANCISCO DE CONCHOS':3030,'SAN FRANCISCO DEL ORO':5027,
        'SANTA BARBARA':12579,'SANTA ISABEL':3814,'SATEVO':3793,'SAUCILLO':29693,
        'TEMOSACHIC':5241,'URIQUE':16988,'URUACHI':7151,'VALLE DE ZARAGOZA':4727,
    }
    import unicodedata as _ud2
    def _nk(s): return _ud2.normalize('NFD',str(s).upper().strip()).encode('ascii','ignore').decode()
    pob_map = {_nk(k): v for k, v in POB_MUN.items()}
    cob_muns = []
    for m in muns:
        pob = pob_map.get(_nk(m['nombre']), 0)
        if pob > 0:
            cob_muns.append({'nombre': m['nombre'], 'benef': m['total'],
                             'pob': pob, 'pct': m['total']/pob*100})
    mun_cob_top = max(cob_muns, key=lambda x: x['pct']) if cob_muns else None
    mun_cob_bot = min((m for m in cob_muns if m['benef'] > 0), key=lambda x: x['pct']) if cob_muns else None

    # Apoyos reales (nivel naranja) — agrupados por nombre sumando instituciones
    apoyos_dict = {}
    for inst_x in insts_ap:
        for ap in inst_x.get('apoyos', []):
            key = ap['nombre']
            if key not in apoyos_dict:
                apoyos_dict[key] = {'nombre': key, 'total': 0, 'mujeres': 0, 'hombres': 0, 'inst': set()}
            apoyos_dict[key]['total']   += ap['total']
            apoyos_dict[key]['mujeres'] += ap['mujeres']
            apoyos_dict[key]['hombres'] += ap['hombres']
            apoyos_dict[key]['inst'].add(inst_x['nombre'])
    progs_ap = [{'nombre': v['nombre'], 'total': v['total'], 'mujeres': v['mujeres'],
                 'hombres': v['hombres'], 'inst': '  ·  '.join(sorted(v['inst']))}
                for v in apoyos_dict.values()]
    progs_ap.sort(key=lambda x: -x['total'])
    ap_top = progs_ap[0] if progs_ap else None
    ap_bot = progs_ap[-1] if progs_ap else None

    add_kpi_table(doc, [
        # Fila 1 — cobertura poblacional
        [{'label': 'Beneficiarios Totales',    'value': fmt(total),
          'sub': 'localizables en el programa'},
         {'label': 'Mujeres Atendidas',         'value': fmt(total_m),
          'sub': pct_of(total_m, total) + ' del total'},
         {'label': 'Hombres Atendidos',         'value': fmt(total_h),
          'sub': pct_of(total_h, total) + ' del total'}],
        # Fila 2 — actividad
        [{'label': 'Municipios Activos',        'value': fmt(mun_activos),
          'sub': f'de 67 — falta(n): {len(muns_sin_cobertura) if muns_sin_cobertura else 0}'},
         {'label': 'Total de Apoyos',           'value': fmt(apoyos_total),
          'sub': str(len(progs_ap)) + ' tipos de apoyo distintos'},
         {'label': 'Instituciones',             'value': str(len(insts)),
          'sub': '  ·  '.join(i['nombre'] for i in insts)}],
        # Fila 3 — rangos de edad
        [{'label': 'Rango Más Atendido',        'value': RLAB[rango_max],
          'sub': fmt(RT[rango_max]) + ' benef. (' + pct_of(RT[rango_max], total) + ')'},
         {'label': 'Rango Menos Atendido',      'value': RLAB[rango_min],
          'sub': fmt(RT[rango_min]) + ' benef. (' + pct_of(RT[rango_min], total) + ')'}],
        # Fila 4 — apoyos más/menos entregados
        [{'label': 'Apoyo Más Entregado',       'value': tc(ap_top['nombre'])[:45] if ap_top else '—',
          'sub': (fmt(ap_top['total']) + ' apoyos  ·  ' + ap_top['inst']) if ap_top else ''},
         {'label': 'Apoyo Menos Entregado',     'value': tc(ap_bot['nombre'])[:45] if ap_bot else '—',
          'sub': (fmt(ap_bot['total']) + ' apoyos  ·  ' + ap_bot['inst']) if ap_bot else ''}],
        # Fila 5 — municipios apoyos
        [{'label': 'Municipio · Más Apoyos',    'value': tc(mun_ap_top['nombre']) if mun_ap_top else '—',
          'sub': fmt(mun_ap_top['total']) + ' apoyos entregados' if mun_ap_top else ''},
         {'label': 'Municipio · Menos Apoyos',  'value': tc(mun_ap_bot['nombre']) if mun_ap_bot else '—',
          'sub': fmt(mun_ap_bot['total']) + ' apoyos entregados' if mun_ap_bot else ''}],
        # Fila 6 — cobertura municipal
        [{'label': 'Mayor Cobertura Municipal',  'value': tc(mun_cob_top['nombre']) if mun_cob_top else '—',
          'sub': f"{mun_cob_top['pct']:.1f}% de su pob. municipal ({fmt(mun_cob_top['benef'])} benef.)" if mun_cob_top else ''},
         {'label': 'Menor Cobertura Municipal',  'value': tc(mun_cob_bot['nombre']) if mun_cob_bot else '—',
          'sub': f"{mun_cob_bot['pct']:.2f}% de su pob. municipal ({fmt(mun_cob_bot['benef'])} benef.)" if mun_cob_bot else ''},
         {'label': 'Municipios sin Cobertura',   'value': str(len(muns_sin_cobertura)) if muns_sin_cobertura else '0',
          'sub': ', '.join(tc(m) for m in muns_sin_cobertura) if muns_sin_cobertura else 'Todos cubiertos'}],
    ])

    # ══ 2. DISTRIBUCIÓN POR SEXO Y RANGO DE EDAD ═════════════════════════════
    add_heading(doc, '2. Distribución por Sexo y Rango de Edad', page_break=True)
    add_body(doc, 'Distribución de los beneficiarios localizables por sexo y rango de edad a nivel estatal.')
    add_spacer(doc, 4)

    # Tabla global sexo
    add_table(doc,
        ['Sexo', 'Beneficiarios', '% del Total'],
        [['Mujeres', fmt(total_m), pct_of(total_m, total)],
         ['Hombres', fmt(total_h), pct_of(total_h, total)],
         ['TOTAL',   fmt(total),   '100.0%']],
        [5320, 2660, 2660]
    )
    add_spacer(doc, 8)

    # Tabla rangos de edad
    add_subheading(doc, 'Rango de edad')
    rows_ed = []
    for r in RANGOS:
        tot_r = RT[r]
        rows_ed.append([RLAB[r],
            fmt(RT_M[r]), pct_of(RT_M[r], tot_r),
            fmt(RT_H[r]), pct_of(RT_H[r], tot_r),
            fmt(tot_r),   pct_of(tot_r, total)])
    rows_ed.append(['TOTAL',
        fmt(total_m), '—', fmt(total_h), '—', fmt(total), '100.0%'])
    add_table(doc,
        ['Rango de Edad', 'Mujeres', '% Muj.', 'Hombres', '% Hom.', 'Total', '% Global'],
        rows_ed, [1600, 1200, 900, 1200, 900, 1200, 900]
    )
    add_spacer(doc, 4)
    p_note = doc.add_paragraph()
    r1 = p_note.add_run('Rango de edad con mayor concentración: ')
    r1.font.name='Arial'; r1.font.size=Pt(9); r1.font.color.rgb=GRIS_TEXT
    r2 = p_note.add_run(f'{RLAB[dom_rango]} años — {fmt(RT[dom_rango])} beneficiarios ({pct_of(RT[dom_rango],total)})')
    r2.font.name='Arial'; r2.font.size=Pt(9); r2.font.bold=True; r2.font.color.rgb=AZUL_GOB
    p_note.paragraph_format.space_before=Pt(4); p_note.paragraph_format.space_after=Pt(2)

    # ══ 3. DISTRIBUCIÓN POR INSTITUCIÓN Y PROGRAMA ═══════════════════════════
    add_heading(doc, '3. Distribución por Institución y Programa', page_break=True)
    add_body(doc, 'Desglose de beneficiarios localizables por institución ejecutora y sus programas activos.')
    add_spacer(doc, 4)

    # Resumen instituciones
    inst_rows = []
    for inst in insts:
        pM_i = sf(inst['mujeres'])/sf(inst['total'])*100 if sf(inst['total']) else 0
        inst_rows.append([
            inst['nombre'], fmt(inst['total']),
            pct_of(inst['total'], total),
            fmt(inst['mujeres']), f'{pM_i:.0f}%',
            fmt(inst['hombres']), f'{100-pM_i:.0f}%',
        ])
    inst_rows.append(['TOTAL', fmt(total), '100.0%',
                      fmt(total_m), pct_of(total_m,total),
                      fmt(total_h), pct_of(total_h,total)])
    add_table(doc,
        ['Institución', 'Beneficiarios', '% Total', 'Mujeres', '% M', 'Hombres', '% H'],
        inst_rows, [2200, 1400, 1000, 1200, 800, 1200, 800],
        first_col_bold_size=9
    )
    add_spacer(doc, 8)

    # Detalle por institución
    for inst in insts:
        add_subheading(doc, tc(inst['nombre']))
        pM_i = sf(inst['mujeres'])/sf(inst['total'])*100 if sf(inst['total']) else 0
        p_desc = doc.add_paragraph()
        r1 = p_desc.add_run(f'Total: '); r1.font.name='Arial'; r1.font.size=Pt(9); r1.font.bold=True; r1.font.color.rgb=AZUL_GOB
        r2 = p_desc.add_run(fmt(inst['total'])+f'  ({pct_of(inst["total"],total)} del programa)   ')
        r2.font.name='Arial'; r2.font.size=Pt(9); r2.font.color.rgb=GRIS_TEXT
        r3 = p_desc.add_run(f'Mujeres: {fmt(inst["mujeres"])} ({pM_i:.0f}%)   Hombres: {fmt(inst["hombres"])} ({100-pM_i:.0f}%)')
        r3.font.name='Arial'; r3.font.size=Pt(9); r3.font.color.rgb=GRIS_TEXT
        p_desc.paragraph_format.space_before=Pt(2); p_desc.paragraph_format.space_after=Pt(4)

        progs_validos = [p for p in inst.get('programas',[]) if p['total'] > 0 and p['nombre'] != inst['nombre']]
        if progs_validos:
            prog_rows = []
            for p in progs_validos:
                pM_p = sf(p['mujeres'])/sf(p['total'])*100 if sf(p['total']) else 0
                prog_rows.append([
                    tc(p['nombre']), fmt(p['total']),
                    pct_of(p['total'], inst['total']),
                    fmt(p['mujeres']), f'{pM_p:.0f}%',
                    fmt(p['hombres']), f'{100-pM_p:.0f}%',
                ])
            add_table(doc,
                ['Programa', 'Beneficiarios', '% Inst.', 'Mujeres', '% M', 'Hombres', '% H'],
                prog_rows, [3500, 1200, 900, 1000, 700, 1000, 700]
            )
        add_spacer(doc, 6)

    # ══ 4. BENEFICIARIOS POR MUNICIPIO ═══════════════════════════════════════
    add_heading(doc, '4. Beneficiarios por Municipio', page_break=True)
    add_body(doc, f'Distribución de los {fmt(total)} beneficiarios localizables entre los {mun_activos} de 67 municipios del estado con cobertura activa en el programa.')
    add_spacer(doc, 4)

    muns_sorted = sorted([m for m in muns if m['total']>0], key=lambda x: -x['total'])
    mun_rows = []
    for i, m in enumerate(muns_sorted):
        pM_m = sf(m['mujeres'])/sf(m['total'])*100 if sf(m['total']) else 0
        mun_rows.append([
            tc(m['nombre']), fmt(m['total']),
            pct_of(m['total'], total),
            fmt(m['mujeres']), f'{pM_m:.0f}%',
            fmt(m['hombres']), f'{100-pM_m:.0f}%',
        ])
    mun_rows.append(['TOTAL', fmt(total), '100.0%',
                     fmt(total_m), pct_of(total_m,total),
                     fmt(total_h), pct_of(total_h,total)])
    add_table(doc,
        ['Municipio', 'Beneficiarios', '% Global', 'Mujeres', '% M', 'Hombres', '% H'],
        mun_rows, [2800, 1400, 1000, 1200, 800, 1200, 800]
    )

    # ══ 5. DETALLE POR RANGO DE EDAD — TOP 10 MUNICIPIOS ══════════════════════
    add_heading(doc, '5. Detalle por Rango de Edad — Top 10 Municipios', page_break=True)
    add_body(doc, 'Distribución por rango de edad de los diez municipios con mayor número de beneficiarios localizables.', keep_next=True)
    add_spacer(doc, 4)

    top10 = muns_sorted[:10]
    headers_t = ['Municipio'] + [RLAB[r] for r in RANGOS] + ['Total']
    rows_t = []
    for m in top10:
        row = [tc(m['nombre'])]
        for r in RANGOS:
            row.append(fmt(m['rangos_m'].get(r,0)+m['rangos_h'].get(r,0)))
        row.append(fmt(m['total']))
        rows_t.append(row)
    # Totals row
    totals_row = ['TOTAL']
    for r in RANGOS: totals_row.append(fmt(RT[r]))
    totals_row.append(fmt(total))
    rows_t.append(totals_row)

    add_table(doc, headers_t, rows_t,
        [2000] + [900]*7 + [900]
    )

    # ══ 6. APOYOS POR INSTITUCIÓN Y PROGRAMA ════════════════════════════════
    add_heading(doc, '6. Apoyos Entregados por Institución y Programa', page_break=True)
    add_body(doc, (
        f'Conteo de apoyos entregados (recuento de registros) por institución y programa. '
        f'Un mismo beneficiario puede recibir múltiples apoyos; el total de apoyos '
        f'({fmt(apoyos_total)}) es mayor que el de beneficiarios únicos ({fmt(total)}).'
    ))
    add_spacer(doc, 6)

    # Tabla tipos de apoyo
    add_subheading(doc, 'Tipos de Apoyo')
    tipos_rows = []
    for p in progs_ap:
        pM_p = sf(p['mujeres'])/sf(p['total'])*100 if sf(p['total']) else 0
        tipos_rows.append([
            tc(p['nombre']), p['inst'],
            fmt(p['total']), pct_of(p['total'], apoyos_total),
            fmt(p['mujeres']), f'{pM_p:.0f}%',
            fmt(p['hombres']), f'{100-pM_p:.0f}%',
        ])
    add_table(doc,
        ['Tipo de Apoyo', 'Institución', 'Apoyos', '% Total', 'Mujeres', '% M', 'Hombres', '% H'],
        tipos_rows, [3000, 1000, 1000, 800, 900, 600, 900, 600]
    )
    add_spacer(doc, 8)

    # Resumen apoyos por institución
    ap_rows = []
    for inst in insts_ap:
        pM_i = sf(inst['mujeres'])/sf(inst['total'])*100 if sf(inst['total']) else 0
        ap_rows.append([
            inst['nombre'], fmt(inst['total']),
            pct_of(inst['total'], apoyos_total),
            fmt(inst['mujeres']), f'{pM_i:.0f}%',
            fmt(inst['hombres']), f'{100-pM_i:.0f}%',
        ])
    ap_rows.append(['TOTAL', fmt(apoyos_total), '100.0%',
                    fmt(total_m_ap), pct_of(total_m_ap, apoyos_total),
                    fmt(total_h_ap), pct_of(total_h_ap, apoyos_total)])
    add_table(doc,
        ['Institución', 'Apoyos', '% Total', 'Mujeres', '% M', 'Hombres', '% H'],
        ap_rows, [2200, 1400, 1000, 1200, 800, 1200, 800],
        first_col_bold_size=9
    )
    add_spacer(doc, 8)

    # Detalle por institución — apoyos
    for inst in insts_ap:
        add_subheading(doc, tc(inst['nombre']))
        pM_i = sf(inst['mujeres'])/sf(inst['total'])*100 if sf(inst['total']) else 0
        p_desc = doc.add_paragraph()
        r1 = p_desc.add_run('Apoyos entregados: '); r1.font.name='Arial'; r1.font.size=Pt(9); r1.font.bold=True; r1.font.color.rgb=AZUL_GOB
        r2 = p_desc.add_run(fmt(inst['total'])+f'  ({pct_of(inst["total"],apoyos_total)} del total)   ')
        r2.font.name='Arial'; r2.font.size=Pt(9); r2.font.color.rgb=GRIS_TEXT
        r3 = p_desc.add_run(f'Mujeres: {fmt(inst["mujeres"])} ({pM_i:.0f}%)   Hombres: {fmt(inst["hombres"])} ({100-pM_i:.0f}%)')
        r3.font.name='Arial'; r3.font.size=Pt(9); r3.font.color.rgb=GRIS_TEXT
        p_desc.paragraph_format.space_before=Pt(2); p_desc.paragraph_format.space_after=Pt(4)

        progs_v = [p for p in inst.get('programas',[]) if p['total'] > 0 and p['nombre'] != inst['nombre']]
        if progs_v:
            prog_rows = []
            for p in progs_v:
                pM_p = sf(p['mujeres'])/sf(p['total'])*100 if sf(p['total']) else 0
                prog_rows.append([
                    tc(p['nombre']), fmt(p['total']),
                    pct_of(p['total'], inst['total']),
                    fmt(p['mujeres']), f'{pM_p:.0f}%',
                    fmt(p['hombres']), f'{100-pM_p:.0f}%',
                ])
            add_table(doc,
                ['Programa', 'Apoyos', '% Inst.', 'Mujeres', '% M', 'Hombres', '% H'],
                prog_rows, [3500, 1200, 900, 1000, 700, 1000, 700]
            )
        add_spacer(doc, 6)

    # ══ 7. APOYOS POR MUNICIPIO ═══════════════════════════════════════════════
    add_heading(doc, '7. Apoyos Entregados por Municipio', page_break=True)
    add_body(doc, f'Distribución de los {fmt(apoyos_total)} apoyos entregados entre los municipios del estado.')
    add_spacer(doc, 4)

    muns_ap_sorted = sorted(muns_ap, key=lambda x: -x['total'])
    mun_ap_rows = []
    for i, m in enumerate(muns_ap_sorted):
        pM_m = sf(m['mujeres'])/sf(m['total'])*100 if sf(m['total']) else 0
        mun_ap_rows.append([
            tc(m['nombre']), fmt(m['total']),
            pct_of(m['total'], apoyos_total),
            fmt(m['mujeres']), f'{pM_m:.0f}%',
            fmt(m['hombres']), f'{100-pM_m:.0f}%',
        ])
    mun_ap_rows.append(['TOTAL', fmt(apoyos_total), '100.0%',
                        fmt(total_m_ap), pct_of(total_m_ap, apoyos_total),
                        fmt(total_h_ap), pct_of(total_h_ap, apoyos_total)])
    add_table(doc,
        ['Municipio', 'Apoyos', '% Global', 'Mujeres', '% M', 'Hombres', '% H'],
        mun_ap_rows, [2800, 1400, 1000, 1200, 800, 1200, 800]
    )

    # ══ 8. CONCLUSIONES Y OBSERVACIONES ══════════════════════════════════════
    add_heading(doc, '8. Conclusiones y Observaciones', page_break=True)

    # Calcular indicadores para las conclusiones
    prom_ap_benef = sf(apoyos_total) / sf(total) if total else 0
    mun_top = muns_sorted[0] if muns_sorted else None
    inst_top = insts[0] if insts else None
    pM_global = sf(total_m)/sf(total)*100 if total else 0

    add_subheading(doc, 'Alcance del Programa')
    add_body(doc, (
        f'Al corte de {MES} {ANO}, el programa NutriChihuahua registra {fmt(total)} beneficiarios '
        f'localizables distribuidos en {mun_activos} municipios del estado de Chihuahua, '
        f'con un total de {fmt(apoyos_total)} apoyos entregados, lo que representa un promedio de '
        f'{prom_ap_benef:.1f} apoyos por beneficiario.'
    ))
    add_spacer(doc, 4)

    add_subheading(doc, 'Composición por Sexo')
    add_body(doc, (
        f'El {pct_of(total_m, total)} de los beneficiarios son mujeres ({fmt(total_m)} personas) '
        f'y el {pct_of(total_h, total)} son hombres ({fmt(total_h)} personas). '
        f'Esta distribución refleja la orientación del programa hacia núcleos familiares '
        f'donde la mujer es el principal receptor del apoyo alimentario.'
    ))
    add_spacer(doc, 4)

    add_subheading(doc, 'Distribución por Edad')
    add_body(doc, (
        f'El rango de edad con mayor concentración de beneficiarios es {RLAB[dom_rango]} años, '
        f'con {fmt(RT[dom_rango])} personas ({pct_of(RT[dom_rango], total)} del total). '
        f'La población de 65 años o más y adultos de 30 a 64 años concentran la mayor proporción '
        f'de apoyos, lo que sugiere un perfil de beneficiario de edad avanzada o adulto en situación '
        f'de vulnerabilidad alimentaria.'
    ))
    add_spacer(doc, 4)

    add_subheading(doc, 'Cobertura Municipal')
    if mun_top:
        add_body(doc, (
            f'El municipio con mayor número de beneficiarios es {tc(mun_top["nombre"])} '
            f'con {fmt(mun_top["total"])} personas ({pct_of(mun_top["total"], total)} del total estatal). '
            f'Los cinco municipios con mayor cobertura concentran el '
            f'{pct_of(sum(m["total"] for m in muns_sorted[:5]), total)} de los beneficiarios del programa.'
        ))
    add_spacer(doc, 4)
    if muns_sin_cobertura:
        p_sin = doc.add_paragraph()
        r1 = p_sin.add_run('Municipio sin cobertura registrada: ')
        r1.font.name='Arial'; r1.font.size=Pt(9); r1.font.bold=True; r1.font.color.rgb=AZUL_GOB
        r2 = p_sin.add_run(', '.join(tc(m) for m in muns_sin_cobertura))
        r2.font.name='Arial'; r2.font.size=Pt(9); r2.font.color.rgb=GRIS_TEXT
        p_sin.paragraph_format.space_before=Pt(2); p_sin.paragraph_format.space_after=Pt(4)
        add_body(doc, (
            f'De los 67 municipios del estado, {len(muns_sin_cobertura)} '
            f'({", ".join(tc(m) for m in muns_sin_cobertura)}) '
            f'no registra beneficiarios en NutriChihuahua al corte de {MES} {ANO}. '
            f'Se recomienda revisar la cobertura en este municipio para garantizar '
            f'la universalidad del programa.'
        ))
    else:
        add_body(doc, f'El programa registra cobertura en los 67 municipios del estado de Chihuahua.')
    add_spacer(doc, 4)

    add_subheading(doc, 'Participación Institucional')
    if inst_top:
        add_body(doc, (
            f'{tc(inst_top["nombre"])} es la institución con mayor participación en el programa, '
            f'atendiendo a {fmt(inst_top["total"])} beneficiarios '
            f'({pct_of(inst_top["total"], total)} del total) y entregando '
            f'{fmt(insts_ap[0]["total"] if insts_ap else 0)} apoyos. '
            f'En total participan {len(insts)} instituciones en la operación del programa: '
            f'{", ".join(tc(i["nombre"]) for i in insts)}.'
        ))
    add_spacer(doc, 8)

    add_subheading(doc, 'Observaciones')
    add_body(doc, (
        '• Los datos presentados corresponden a beneficiarios localizables registrados en el padrón; '
        'el número real de personas atendidas puede ser mayor al incluir beneficiarios no localizables.'
    ))
    add_spacer(doc, 2)
    add_body(doc, (
        '• El conteo de apoyos es mayor al de beneficiarios únicos porque un mismo beneficiario '
        'puede recibir apoyos de distintos programas o instituciones de manera simultánea.'
    ))
    add_spacer(doc, 2)
    add_body(doc, (
        f'• La información fue generada a partir del archivo de padrón de beneficiarios '
        f'al corte de {MES} {ANO} y refleja el estado del registro en esa fecha.'
    ))
    add_spacer(doc, 8)

    # ══ PIE DE DOCUMENTO ══════════════════════════════════════════════════════
    add_spacer(doc, 20)
    p_fin = doc.add_paragraph()
    p_fin.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_top_border(p_fin)
    r_fin = p_fin.add_run(f'Secretaría de Desarrollo Humano y Bien Común  ·  Gobierno del Estado de Chihuahua  ·  {ANO}')
    r_fin.font.name='Arial'; r_fin.font.size=Pt(8); r_fin.font.italic=True
    r_fin.font.color.rgb=RGBColor(0x88,0x88,0x88)

    out = Path(OUT_PATH)
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))
    kb = out.stat().st_size // 1024
    print(f'Documento guardado: {out.name}  ({kb} KB)')

if __name__ == '__main__':
    main()
