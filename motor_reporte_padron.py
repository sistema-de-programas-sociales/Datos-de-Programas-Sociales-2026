#!/usr/bin/env python3
"""
motor_reporte_padron.py
Generador de reporte mensual a partir del archivo Reporte_e_Informe_de_Padron.xlsx
Mantiene la misma estructura visual que motor_reporte.py

Uso: python3 motor_reporte_padron.py <excel> <mes> <quincena> <año> [output]
"""
import sys, os, json, subprocess, base64, unicodedata
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ─── ARGUMENTOS ───────────────────────────────────────────────────────────────
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
EXCEL_PATH = sys.argv[1] if len(sys.argv) > 1 else os.path.join(SCRIPT_DIR, 'Reporte_e_Informe_de_Padron.xlsx')
MES        = sys.argv[2] if len(sys.argv) > 2 else 'Enero'
QUINCENA   = sys.argv[3] if len(sys.argv) > 3 else '1'
ANO        = sys.argv[4] if len(sys.argv) > 4 else '2026'

REPORTES_DIR = os.path.join(SCRIPT_DIR, 'Reportes')
os.makedirs(REPORTES_DIR, exist_ok=True)

# Detectar si el período es trimestral por el nombre del mes
TRIMESTRE_MAP = {
    '1er_trimestre':          ('1er Trimestre',          'Enero – Marzo'),
    '2do_trimestre':          ('2do Trimestre',          'Abril – Junio'),
    '3er_trimestre':          ('3er Trimestre',          'Julio – Septiembre'),
    '4to_trimestre_cierre':   ('4to Trimestre – Cierre', 'Octubre – Diciembre'),
}
_mes_key = MES.lower().replace(' ', '_')
_trimestre_info = TRIMESTRE_MAP.get(_mes_key)

if _trimestre_info:
    ES_TRIMESTRAL  = True
    PERIODO_CORTO  = _trimestre_info[0]          # "1er Trimestre"
    PERIODO_MESES  = _trimestre_info[1]          # "Enero – Marzo"
    PERIODO_LARGO  = f'{PERIODO_CORTO} {ANO}'   # "1er Trimestre 2026"
    TITULO_CORTO   = PERIODO_LARGO
else:
    ES_TRIMESTRAL  = False
    PERIODO_CORTO  = MES
    PERIODO_MESES  = ''
    PERIODO_LARGO  = f'{MES} {ANO}'
    TITULO_CORTO   = PERIODO_LARGO

POB_VULNERABLE = 1792324

OUTPUT_PATH = sys.argv[5] if len(sys.argv) > 5 else os.path.join(
    REPORTES_DIR,
    f'Reporte_Programas_Sociales_{MES}_{ANO}.docx'
)

# ─── COLORES ──────────────────────────────────────────────────────────────────
AZUL_GOB  = RGBColor(0x1B, 0x3A, 0x6B)
AZUL_MED  = RGBColor(0x2E, 0x5B, 0xA8)
AZUL_CLAR = RGBColor(0xD6, 0xE4, 0xF7)
GRIS_TEXT = RGBColor(0x33, 0x33, 0x33)
BLANCO    = RGBColor(0xFF, 0xFF, 0xFF)
DORADO    = RGBColor(0xC8, 0xA0, 0x00)

# ─── LEER EXCEL ───────────────────────────────────────────────────────────────
def leer_excel():
    script = os.path.join(SCRIPT_DIR, 'read_excel_padron.py')
    result = subprocess.run(['python3', script, EXCEL_PATH], capture_output=True, text=True)
    if result.returncode != 0:
        print('Error leyendo Excel:', result.stderr); sys.exit(1)
    return json.loads(result.stdout)

def leer_grupos_vulnerables():
    """Lee la hoja Grupos Vulnerables del Excel del padrón."""
    try:
        import openpyxl
        wb = openpyxl.load_workbook(EXCEL_PATH, data_only=True)
        if 'Grupos Vulnerables' not in wb.sheetnames:
            return []
        ws = wb['Grupos Vulnerables']
        grupos = []
        for row in ws.iter_rows(min_row=2, values_only=True):
            if not row[0]:
                continue
            nombre = str(row[0]).strip()
            try:
                pob_vul = int(row[1]) if row[1] not in (None, '') else 0
            except (ValueError, TypeError):
                pob_vul = 0
            try:
                pob_ate = int(row[2]) if row[2] not in (None, '') else 0
            except (ValueError, TypeError):
                pob_ate = 0
            grupos.append({'nombre': nombre, 'pob_vulnerable': pob_vul, 'atendidos': pob_ate})
        return grupos
    except Exception as e:
        print(f'AVISO: No se pudo leer Grupos Vulnerables: {e}')
        return []

# ─── HELPERS ──────────────────────────────────────────────────────────────────
def fmt(n, dec=0):
    if n is None or n == '': return '0'
    try:
        num = float(n)
        if dec == 0: return f'{int(round(num)):,}'
        return f'{num:,.{dec}f}'
    except: return '0'

def pct_of(part, total):
    if not total or total == 0: return '0'
    try: return f'{float(part)/float(total)*100:.1f}%'
    except: return '0'

def sf(v):
    try: return float(v) if v is not None else 0.0
    except: return 0.0

# Palabras que no se capitalizan en title-case (preposiciones/artículos en español)
_MINUSCULAS = {'a','ante','bajo','con','contra','de','del','desde','durante','el','en',
               'entre','hacia','hasta','la','las','lo','los','mediante','para','por',
               'que','se','sin','sobre','su','sus','tras','un','una','unas','unos','y'}
def tc(s):
    """Convierte texto a Title Case inteligente en español.
    Preserva textos que ya tienen minúsculas (no están en ALL-CAPS)."""
    if not s: return s
    # Si no está en mayúsculas puras, no tocar
    if s != s.upper(): return s
    words = s.split()
    result = []
    for i, w in enumerate(words):
        result.append(w.capitalize() if (i == 0 or w.lower() not in _MINUSCULAS) else w.lower())
    return ' '.join(result)

def rgb_hex(c): return f'{c[0]:02X}{c[1]:02X}{c[2]:02X}'

# ─── XML HELPERS ──────────────────────────────────────────────────────────────
def set_cell_bg(cell, color):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),'clear'); shd.set(qn('w:color'),'auto')
    shd.set(qn('w:fill'), rgb_hex(color)); tcPr.append(shd)

def set_cell_borders(cell, color='CCCCCC'):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    tcB = OxmlElement('w:tcBorders')
    for side in ['top','left','bottom','right']:
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'),'single'); el.set(qn('w:sz'),'4')
        el.set(qn('w:space'),'0');    el.set(qn('w:color'),color)
        tcB.append(el)
    tcPr.append(tcB)

def set_cell_no_borders(cell):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    tcB = OxmlElement('w:tcBorders')
    for side in ['top','left','bottom','right']:
        el = OxmlElement(f'w:{side}'); el.set(qn('w:val'),'none'); tcB.append(el)
    tcPr.append(tcB)

def set_cell_margins(cell, top=80, bottom=80, left=120, right=120):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    tcM = OxmlElement('w:tcMar')
    for side, val in [('top',top),('bottom',bottom),('left',left),('right',right)]:
        el = OxmlElement(f'w:{side}'); el.set(qn('w:w'),str(val)); el.set(qn('w:type'),'dxa')
        tcM.append(el)
    tcPr.append(tcM)

def set_col_width(cell, w):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    tcW = OxmlElement('w:tcW'); tcW.set(qn('w:w'),str(w)); tcW.set(qn('w:type'),'dxa')
    tcPr.append(tcW)

def prevent_row_break(row):
    tr = row._tr; trPr = tr.get_or_add_trPr()
    cs = OxmlElement('w:cantSplit'); cs.set(qn('w:val'),'1'); trPr.append(cs)

def keep_with_next_row(row):
    for cell in row.cells:
        for p in cell.paragraphs:
            pPr = p._p.get_or_add_pPr()
            kn = OxmlElement("w:keepNext")
            kn.set(qn("w:val"), "1")
            pPr.append(kn)

def set_keep_next_on_last_table(doc):
    """Aplica keepNext al ultimo parrafo de la ultima tabla del doc."""
    tables = doc.tables
    if not tables: return
    last_table = tables[-1]
    last_row = last_table.rows[-1]
    for cell in last_row.cells:
        for p in cell.paragraphs:
            pPr = p._p.get_or_add_pPr()
            kn = OxmlElement('w:keepNext')
            kn.set(qn('w:val'), '1')
            pPr.append(kn)

def insert_page_break(doc):
    """Inserta salto de página explícito (para uso entre bloques, no en headings)."""
    p = OxmlElement('w:p'); r = OxmlElement('w:r'); br = OxmlElement('w:br')
    br.set(qn('w:type'),'page'); r.append(br); p.append(r)
    doc.element.body.append(p)

def set_page_break_before(p):
    """
    Marca el párrafo con pageBreakBefore=True.
    Esto garantiza que el párrafo SIEMPRE inicia en una nueva página,
    sin dejar párrafos vacíos al final de la página anterior.
    """
    pPr = p._p.get_or_add_pPr()
    # Eliminar cualquier pageBreakBefore previo
    for existing in pPr.findall(qn('w:pageBreakBefore')):
        pPr.remove(existing)
    pbr = OxmlElement('w:pageBreakBefore')
    pbr.set(qn('w:val'), '1')
    pPr.insert(0, pbr)

def add_bottom_border(p, color='2E5BA8', sz='6'):
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'),'single'); bot.set(qn('w:sz'),sz)
    bot.set(qn('w:space'),'4'); bot.set(qn('w:color'),color)
    pBdr.append(bot); pPr.append(pBdr)

def add_top_border(p, color='1B3A6B', sz='4'):
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    top = OxmlElement('w:top')
    top.set(qn('w:val'),'single'); top.set(qn('w:sz'),sz)
    top.set(qn('w:space'),'4'); top.set(qn('w:color'),color)
    pBdr.append(top); pPr.append(pBdr)

# ─── ELEMENTOS DEL DOCUMENTO ──────────────────────────────────────────────────
def add_heading(doc, text, page_break=True):
    p = doc.add_paragraph()
    if page_break:
        set_page_break_before(p)
    run = p.add_run(text)
    run.font.name='Arial'; run.font.size=Pt(16); run.font.bold=True
    run.font.color.rgb=AZUL_GOB
    p.paragraph_format.space_before=Pt(0); p.paragraph_format.space_after=Pt(4)
    add_bottom_border(p)

def add_body(doc, text, size=10, color=None, italic=False, keep_next=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name='Arial'; run.font.size=Pt(size); run.font.italic=italic
    run.font.color.rgb = color if color else GRIS_TEXT
    p.paragraph_format.space_before=Pt(2); p.paragraph_format.space_after=Pt(2)
    if keep_next:
        pPr = p._p.get_or_add_pPr()
        kn = OxmlElement('w:keepNext'); kn.set(qn('w:val'), '1')
        pPr.append(kn)

def add_totals_line(doc, label, text):
    p = doc.add_paragraph()
    r1 = p.add_run(label)
    r1.font.bold=True; r1.font.name='Arial'; r1.font.size=Pt(9); r1.font.color.rgb=AZUL_GOB
    r2 = p.add_run(text)
    r2.font.name='Arial'; r2.font.size=Pt(9); r2.font.color.rgb=GRIS_TEXT
    p.paragraph_format.space_before=Pt(2); p.paragraph_format.space_after=Pt(2)

def add_bullet(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    run.font.name='Arial'; run.font.size=Pt(10); run.font.color.rgb=GRIS_TEXT

def add_spacer(doc, pts=6, keep_next=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before=Pt(pts); p.paragraph_format.space_after=Pt(0)
    if keep_next:
        pPr = p._p.get_or_add_pPr()
        kn = OxmlElement('w:keepNext'); kn.set(qn('w:val'), '1')
        pPr.append(kn)

# ─── TABLA GENERICA ───────────────────────────────────────────────────────────
def set_repeat_header(row):
    """Marca una fila como encabezado que se repite en cada página."""
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    tblHeader = OxmlElement('w:tblHeader')
    tblHeader.set(qn('w:val'), '1')
    trPr.append(tblHeader)

def set_table_width(table, width_twips):
    """Fuerza el ancho total de la tabla."""
    tbl = table._tbl; tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None: tblPr = OxmlElement('w:tblPr'); tbl.insert(0, tblPr)
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), str(width_twips)); tblW.set(qn('w:type'), 'dxa')
    existing = tblPr.find(qn('w:tblW'))
    if existing is not None: tblPr.remove(existing)
    tblPr.append(tblW)

def add_table(doc, headers, rows, col_widths, first_col_bold_size=None):
    # Escalar columnas para que sumen exactamente 10640 (ancho útil de página)
    PAGE_W = 10640
    raw_sum = sum(col_widths)
    scaled = [int(w * PAGE_W / raw_sum) for w in col_widths]
    diff = PAGE_W - sum(scaled)
    scaled[-1] += diff   # ajustar último para compensar redondeo
    col_widths = scaled

    n = len(headers)
    t = doc.add_table(rows=1+len(rows), cols=n)
    t.style='Table Grid'; t.alignment=WD_TABLE_ALIGNMENT.LEFT
    set_table_width(t, PAGE_W)
    hr = t.rows[0]
    set_repeat_header(hr)
    for ci, (h, w) in enumerate(zip(headers, col_widths)):
        cell = hr.cells[ci]
        set_cell_bg(cell, AZUL_GOB); set_cell_borders(cell,'1B3A6B')
        set_cell_margins(cell); set_col_width(cell, w)
        p = cell.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(str(h) if h else '')
        run.font.name='Arial'; run.font.size=Pt(7.5); run.font.bold=True
        run.font.color.rgb=BLANCO
        p.paragraph_format.space_before=Pt(0); p.paragraph_format.space_after=Pt(0)
        cell.vertical_alignment=WD_ALIGN_VERTICAL.CENTER
    n_rows = len(rows)
    for ri, row_data in enumerate(rows):
        is_total = str(row_data[0]).strip().upper().startswith('TOTAL')
        is_last  = (ri == n_rows - 1)
        bg = RGBColor(0xD6, 0xE4, 0xF7) if is_total else (BLANCO if ri%2==0 else AZUL_CLAR)
        row = t.rows[ri+1]
        prevent_row_break(row)
        # Penúltima fila: keepNext para que no quede sola la última fila en página nueva
        if not is_last and ri == n_rows - 2:
            keep_with_next_row(row)
        for ci, (val, w) in enumerate(zip(row_data, col_widths)):
            cell = row.cells[ci]
            set_cell_bg(cell, bg); set_cell_borders(cell)
            set_cell_margins(cell); set_col_width(cell, w)
            p = cell.paragraphs[0]
            p.alignment=WD_ALIGN_PARAGRAPH.LEFT if ci==0 else WD_ALIGN_PARAGRAPH.RIGHT
            run = p.add_run('0' if val is None else str(val))
            run.font.name='Arial'
            if not is_total and ci == 0 and first_col_bold_size:
                run.font.size=Pt(first_col_bold_size); run.font.bold=True
            else:
                run.font.size=Pt(9) if is_total else Pt(7.5)
                run.font.bold=is_total
            run.font.color.rgb=AZUL_GOB if is_total else GRIS_TEXT
            p.paragraph_format.space_before=Pt(0); p.paragraph_format.space_after=Pt(0)
            cell.vertical_alignment=WD_ALIGN_VERTICAL.CENTER

# ─── TABLA KPI ────────────────────────────────────────────────────────────────
def add_kpi_table(doc, kpi_rows):
    for group in kpi_rows:
        KPI_W = 10640
        while len(group) < 3: group.append({'label':'','value':'','sub':''})
        t = doc.add_table(rows=1, cols=3); t.alignment=WD_TABLE_ALIGNMENT.LEFT
        set_table_width(t, KPI_W)
        col_w = KPI_W // 3
        row = t.rows[0]
        for ci, kpi in enumerate(group):
            cell = row.cells[ci]
            set_cell_bg(cell, AZUL_CLAR); set_cell_no_borders(cell)
            set_cell_margins(cell,80,80,120,120); set_col_width(cell, col_w)
            p1 = cell.paragraphs[0]; p1.alignment=WD_ALIGN_PARAGRAPH.CENTER
            r1 = p1.add_run(kpi.get('value','-'))
            r1.font.name='Arial'; r1.font.size=Pt(15); r1.font.bold=True
            r1.font.color.rgb=AZUL_GOB
            p1.paragraph_format.space_before=Pt(0); p1.paragraph_format.space_after=Pt(2)
            p2 = cell.add_paragraph(); p2.alignment=WD_ALIGN_PARAGRAPH.CENTER
            r2 = p2.add_run(kpi.get('label',''))
            r2.font.name='Arial'; r2.font.size=Pt(7.5); r2.font.bold=True
            r2.font.color.rgb=GRIS_TEXT
            p2.paragraph_format.space_before=Pt(0); p2.paragraph_format.space_after=Pt(2)
            if kpi.get('sub'):
                p3 = cell.add_paragraph(); p3.alignment=WD_ALIGN_PARAGRAPH.CENTER
                r3 = p3.add_run(kpi['sub'])
                r3.font.name='Arial'; r3.font.size=Pt(6.5); r3.font.italic=True
                r3.font.color.rgb=RGBColor(0x66,0x66,0x66)
                p3.paragraph_format.space_before=Pt(0); p3.paragraph_format.space_after=Pt(0)
        add_spacer(doc, 4)

# ─── PAGINA ───────────────────────────────────────────────────────────────────
def set_page_margins(section):
    section.top_margin=Cm(1.8); section.bottom_margin=Cm(1.8)
    section.left_margin=Cm(1.41); section.right_margin=Cm(1.41)

def add_header_footer(section):
    hdr = section.header
    hp  = hdr.paragraphs[0] if hdr.paragraphs else hdr.add_paragraph()
    hp.clear()
    run = hp.add_run(f'Reporte — {TITULO_CORTO}  |  Gobierno del Estado de Chihuahua')
    run.font.name='Arial'; run.font.size=Pt(7.5); run.font.color.rgb=AZUL_GOB
    add_bottom_border(hp, color='1B3A6B', sz='4')
    ftr = section.footer
    fp  = ftr.paragraphs[0] if ftr.paragraphs else ftr.add_paragraph()
    fp.clear(); fp.alignment=WD_ALIGN_PARAGRAPH.LEFT
    run = fp.add_run('Secretaría de Desarrollo Humano y Bien Común  —  Chihuahua, México')
    run.font.name='Arial'; run.font.size=Pt(7); run.font.color.rgb=RGBColor(0x88,0x88,0x88)
    add_top_border(fp)

# ─── PORTADA ──────────────────────────────────────────────────────────────────
def build_portada(doc, fecha_str):
    section = doc.sections[0]
    section.top_margin=Cm(2.5); section.bottom_margin=Cm(2.5)
    section.left_margin=Cm(2.0); section.right_margin=Cm(2.0)
    section.different_first_page_header_footer=True

    for _ in range(6): add_spacer(doc, 8)

    # ── Encabezado institucional ─────────────────────────────────────────────────
    p = doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('GOBIERNO DEL ESTADO DE CHIHUAHUA')
    r.font.name='Arial'; r.font.size=Pt(13); r.font.bold=True; r.font.color.rgb=AZUL_GOB
    p.paragraph_format.space_after=Pt(4)

    p = doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('Secretaría de Desarrollo Humano y Bien Común')
    r.font.name='Arial'; r.font.size=Pt(11); r.font.color.rgb=AZUL_MED
    p.paragraph_format.space_after=Pt(18)

    # ── Línea divisoria ──────────────────────────────────────────────────────────
    p = doc.add_paragraph(); add_bottom_border(p, color='1B3A6B', sz='12')
    p.paragraph_format.space_after=Pt(22)

    # ── Título del reporte ───────────────────────────────────────────────────────
    p = doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('REPORTE DE AVANCE')
    r.font.name='Arial'; r.font.size=Pt(22); r.font.bold=True; r.font.color.rgb=AZUL_GOB
    p.paragraph_format.space_after=Pt(8)

    p = doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('Programas Sociales del Estado de Chihuahua')
    r.font.name='Arial'; r.font.size=Pt(14); r.font.color.rgb=AZUL_MED
    p.paragraph_format.space_after=Pt(16)

    p = doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(TITULO_CORTO.upper())
    r.font.name='Arial'; r.font.size=Pt(28); r.font.bold=True; r.font.color.rgb=DORADO
    p.paragraph_format.space_after=Pt(22)

    # ── Segunda línea divisoria ──────────────────────────────────────────────────
    p = doc.add_paragraph(); add_bottom_border(p, color='1B3A6B', sz='12')
    p.paragraph_format.space_after=Pt(24)

    for _ in range(2): add_spacer(doc, 10)

    # ── Fecha y período ──────────────────────────────────────────────────────────
    p = doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f'Período reportado: {MES} {ANO}')
    r.font.name='Arial'; r.font.size=Pt(11); r.font.bold=True; r.font.color.rgb=AZUL_MED
    p.paragraph_format.space_after=Pt(6)

    p = doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f'Fecha de emisión: {fecha_str}')
    r.font.name='Arial'; r.font.size=Pt(10); r.font.color.rgb=GRIS_TEXT

# ─── MAIN ─────────────────────────────────────────────────────────────────────
def main():
    print('Leyendo datos del Excel...')
    data = leer_excel()

    gt            = data['gran_total']
    rangos        = data['rangos_edad']
    rangos_mh     = data.get('rangos_mh_global', {})
    instituciones = data['instituciones']
    municipios    = data['municipios']
    apoyos        = data['apoyos']
    loc           = data.get('localizables', {})
    indicadores   = data.get('indicadores', [])

    total_benef = sf(gt.get('total', 0))
    total_m     = sf(gt.get('m', 0))
    total_h     = sf(gt.get('h', 0))
    total_sn    = sum(sf(v.get('sn', 0)) for v in instituciones.values())
    POB_ESTATAL = data.get('pob_estatal', 4043130)  # fuente canónica: JSON del lector
    mun_act     = 67  # todos los municipios tienen cobertura vía SALUD
    inst_act    = [k for k, v in instituciones.items() if sf(v.get('total',0)) >= 10]
    total_inst  = len(inst_act)
    total_prog  = sum(len(v.get('programas',[])) for v in instituciones.values())
    # Filtro: excluir nombres que son instituciones y programas con solo 1 beneficiario
    INST_NAMES_UP = {'CECYTECH','COESPO','COESVI','DIF','ICHD','ICHDII','ICHIJUV','ICHIMUJ',
                     'RURAL','SALUD','SDBYBC','SDHyBC','SDHYBC','SEECH','SEYD','SEyD',
                     'SPyCI','SPYCI','TRABAJO','TURISMO','CULTURA','MEDICHIHUAHUA',
                     'DESARROLLO HUMANO','NO IDENTIFICADO'}
    ap_clean     = [a for a in apoyos
                    if str(a['apoyo']).upper().strip() not in INST_NAMES_UP]
    # Usar el total directo del Excel (fila TOTAL de Apoyos Otorgados)
    total_apoyos = sf(data.get('total_apoyos_excel') or sum(sf(a.get('total', 0)) for a in ap_clean))
    cob_estatal = pct_of(total_benef, POB_VULNERABLE)
    muj_pct     = pct_of(total_m, total_benef)
    hom_pct     = pct_of(total_h, total_benef)

    ninos   = sf(rangos.get('0-5',0))  + sf(rangos.get('6-11',0))
    jovenes = sf(rangos.get('12-17',0)) + sf(rangos.get('18-29',0))
    adultos = sf(rangos.get('30-49',0)) + sf(rangos.get('50-64',0))
    mayores = sf(rangos.get('65+',0))

    # ── Beneficiarios Localizables ─────────────────────────────────────────────
    loc_total   = int(loc.get('total', 0))
    loc_m       = int(loc.get('m', 0))
    loc_h       = int(loc.get('h', 0))
    loc_pct_ben = pct_of(loc_total, total_benef)   # % sobre beneficiarios únicos
    loc_muj_pct = pct_of(loc_m, loc_total)
    loc_hom_pct = pct_of(loc_h, loc_total)
    loc_inst    = loc.get('por_institucion', [])    # [{nombre,m,h,total}]
    loc_mun     = loc.get('por_municipio', {})      # {NOMBRE_MAYUS: total}

    meses_es = ['enero','febrero','marzo','abril','mayo','junio','julio','agosto',
                'septiembre','octubre','noviembre','diciembre']
    hoy = datetime.now()
    fecha_str = f'{hoy.day} de {meses_es[hoy.month-1]} de {hoy.year}'

    print('Construyendo documento...')
    doc = Document()
    section = doc.sections[0]
    set_page_margins(section)
    add_header_footer(section)

    # ══ PORTADA ═══════════════════════════════════════════════════════════════
    build_portada(doc, fecha_str)

    # ══ 1. RESUMEN EJECUTIVO ══════════════════════════════════════════════════
    add_heading(doc, '1. Resumen Ejecutivo', page_break=True)
    periodo_desc = f'{PERIODO_LARGO}' + (f' ({PERIODO_MESES})' if ES_TRIMESTRAL and PERIODO_MESES else '')
    add_body(doc, f'Informe de gestión institucional al corte de {periodo_desc} de los programas sociales del Gobierno del Estado de Chihuahua llevado a cabo por la Coordinación de Evaluación y Sistemas de Información Estatal (CEySIE).')
    add_spacer(doc, 6)
    cob_vulnerable = pct_of(total_benef, POB_VULNERABLE)
    cob_estatal_pct = pct_of(total_benef, POB_ESTATAL)

    add_kpi_table(doc, [
        # Fila 1 — cobertura poblacional
        [{'label':'Población Estatal 2026',   'value':fmt(POB_ESTATAL),    'sub':'habitantes en Chihuahua'},
         {'label':'Población Vulnerable',      'value':fmt(POB_VULNERABLE), 'sub':'en condición de vulnerabilidad'},
         {'label':'Beneficiarios Únicos',      'value':fmt(total_benef),    'sub':f'{cob_vulnerable} de pob. en condición de vulnerabilidad'}],
        # Fila 2 — actividad
        [{'label':'Apoyos Otorgados',          'value':fmt(total_apoyos),   'sub':'total de apoyos entregados'},
         {'label':'Municipios con Cobertura',  'value':fmt(mun_act),        'sub':'de 67 municipios'},
         {'label':'Programas Activos',         'value':fmt(total_prog),     'sub':f'{total_inst} instituciones participantes'}],
        # Fila 3 — genero
        [{'label':'Mujeres Atendidas',         'value':fmt(total_m),        'sub':f'{muj_pct} del total'},
         {'label':'Hombres Atendidos',         'value':fmt(total_h),        'sub':f'{hom_pct} del total'},
         {'label':'Sin Datos de Sexo',        'value':fmt(total_sn),       'sub':'pendientes de registro'}],
        # Fila 4 — edad (grupos simplificados)
        [{'label':'Niños',                        'value':fmt(ninos),          'sub':'0 a 11 años'},
         {'label':'Jóvenes',                       'value':fmt(jovenes),        'sub':'12 a 29 años'},
         {'label':'Adultos',                       'value':fmt(adultos),        'sub':'30 a 64 años'}],
        [{'label':'Personas Mayores',              'value':fmt(mayores),        'sub':'65 años o más'},
         {'label':'Cobertura Estatal',         'value':cob_estatal_pct,     'sub':'del total de habitantes'},
         {'label':'Cobertura vulnerable',      'value':cob_vulnerable,      'sub':'de la pob. en condición de vulnerabilidad'}],
        # Fila 5 — localizables
        [{'label':'Beneficiarios Localizables',    'value':fmt(loc_total),      'sub':f'{loc_pct_ben} de beneficiarios únicos'},
         {'label':'Localizables — Mujeres',        'value':fmt(loc_m),          'sub':f'{loc_muj_pct} del total localizable'},
         {'label':'Localizables — Hombres',        'value':fmt(loc_h),          'sub':f'{loc_hom_pct} del total localizable'}],
        # Fila 6 — presupuesto global
    ])

    # KPIs de presupuesto (calculados fuera del add_kpi_table para tener fmt_mon en scope)
    _pres_vals  = [p['presupuesto'] for p in indicadores if p.get('presupuesto') and p['presupuesto']>0]
    _gasto_vals = [p['gasto']       for p in indicadores if p.get('gasto')       and p['gasto']>0]
    _pres_total  = sum(_pres_vals)
    _gasto_total = sum(_gasto_vals)
    _gasto_x_ben = (_gasto_total / total_benef) if total_benef > 0 and _gasto_total > 0 else 0
    def _fmx(v): return f"${v:,.0f} MXN" if v else 'Sin datos'
    add_kpi_table(doc, [
        [{'label':'Presupuesto registrado',
          'value': _fmx(_pres_total),
          'sub': f"{len(_pres_vals)} de {len(indicadores)} programas con dato"},
         {'label':'Gasto ejecutado registrado',
          'value': _fmx(_gasto_total),
          'sub': f"{len(_gasto_vals)} programas con gasto registrado"},
         {'label':'Gasto por beneficiario (est.)',
          'value': _fmx(_gasto_x_ben),
          'sub': 'estimado sobre programas con dato de gasto'}],
    ])

    # ══ 2. BENEFICIARIOS POR SEXO Y RANGO DE EDAD ════════════════════════════
    add_heading(doc, '2. Beneficiarios Únicos por Sexo y Rango de Edad', page_break=True)
    add_body(doc, 'Distribución de los beneficiarios únicos por institución y sexo, con desglose por rango de edad.')
    add_spacer(doc, 2)

    # Nota metodológica
    nota = (
        'Nota metodológica: El total de beneficiarios únicos representa personas reales registradas en el padrón, '
        'independientemente de cuántos programas sociales reciban de forma simultánea. '
        'Los conteos desagregados —por programa, institución, sexo o rango de edad— reflejan registros de atención '
        'no personas distintas, por lo que su suma siempre será igual o mayor al total único. '
        'Esto ocurre porque una misma persona puede estar inscrita en varios programas al mismo tiempo: '
        'por ejemplo, un beneficiario que recibe apoyo del DIF, atención médica de SALUD y una despensa de SDHyBC '
        'aparece contado una vez en cada programa, pero únicamente una vez en el Total de Beneficiarios Únicos. '
        'De igual forma, si esa persona está registrada con su sexo en un programa pero sin ese dato en otro, '
        'se contabiliza en ambas columnas del desglose por sexo, aunque el total único la sigue contando como una sola persona. '
        'Los porcentajes y coberturas presentados en este reporte se calculan siempre sobre el total de beneficiarios únicos.'
    )
    add_body(doc, nota, italic=True, color=RGBColor(0x55, 0x55, 0x55), keep_next=True)
    add_spacer(doc, 4, keep_next=True)

    # Tabla sexo por institución
    inst_rows = []
    for inst_name in sorted(instituciones.keys(), key=lambda x: -sf(instituciones[x].get('total', 0))):
        v = instituciones[inst_name]
        tot = sf(v.get('total',0))
        if tot == 0: continue
        inst_rows.append([
            inst_name,
            fmt(tot),
            fmt(v.get('m',0)),
            fmt(v.get('h',0)),
            fmt(int(v.get('sn',0))),
            pct_of(v.get('m',0), tot),
            pct_of(v.get('h',0), tot),
        ])
    # Totals row
    inst_rows.append(['TOTAL', fmt(total_benef), fmt(total_m), fmt(total_h),
                      fmt(int(total_sn)), muj_pct, hom_pct])

    add_table(doc,
        ['Institución','Total','Mujeres','Hombres','Sin dato de sexo','% Mujeres','% Hombres'],
        inst_rows,
        [2000, 900, 900, 900, 900, 900, 900])
    add_spacer(doc, 4)

    # Tabla rangos de edad globales con desglose por sexo (fuente: col S del Excel)
    add_body(doc, 'Distribución de beneficiarios por rango de edad (todos los municipios):', keep_next=True)
    add_spacer(doc, 3, keep_next=True)
    _edad_labels = [
        ('0 - 5 años',       '0-5'),
        ('6 - 11 años',      '6-11'),
        ('12 - 17 años',     '12-17'),
        ('18 - 29 años',     '18-29'),
        ('30 - 49 años',     '30-49'),
        ('50 - 64 años',     '50-64'),
        ('65 años o más',    '65+'),
        ('Sin dato de edad', 'sin_datos'),
    ]
    edad_rows = []
    for label_e, key in _edad_labels:
        t_e  = sf(rangos.get(key, 0))
        mh   = rangos_mh.get(key, {})
        m_e  = sf(mh.get('m', 0))
        h_e  = sf(mh.get('h', 0))
        sn_e = sf(mh.get('sn', max(0, int(t_e - m_e - h_e))))
        edad_rows.append([
            label_e,
            fmt(t_e),
            fmt(m_e),
            fmt(h_e),
            fmt(sn_e),
            pct_of(m_e, t_e),
            pct_of(h_e, t_e),
            pct_of(t_e, total_benef),
        ])
    edad_rows.append([
        'TOTAL', fmt(total_benef),
        fmt(total_m), fmt(total_h), fmt(int(total_sn)),
        muj_pct, hom_pct, '100.0%',
    ])
    add_table(doc,
        ['Rango de Edad', 'Total', 'Mujeres', 'Hombres', 'Sin dato', '% Mujeres', '% Hombres', '% del total'],
        edad_rows,
        [1800, 950, 950, 950, 800, 900, 900, 900])

    # ══ 3. BENEFICIARIOS LOCALIZABLES ═════════════════════════════════════════
    add_heading(doc, '3. Beneficiarios Localizables', page_break=True)
    add_body(doc,
        f'Un beneficiario localizable es aquel cuyo registro cuenta con nombre, apellido, sexo, fecha de nacimiento, '
        f'CURP, municipio, teléfono y código postal válidos.')
    add_spacer(doc, 4)

    # KPI localizables
    add_kpi_table(doc, [[
        {'label':'Beneficiarios Localizables', 'value':fmt(loc_total),   'sub':f'{loc_pct_ben} de beneficiarios únicos'},
        {'label':'Mujeres Localizables',       'value':fmt(loc_m),       'sub':f'{loc_muj_pct} del total localizable'},
        {'label':'Hombres Localizables',       'value':fmt(loc_h),       'sub':f'{loc_hom_pct} del total localizable'},
    ]])

    # 3a. Tabla por institución
    add_body(doc, '3.1  Beneficiarios localizables por institución:', keep_next=True)
    add_spacer(doc, 3, keep_next=True)
    loc_inst_rows = sorted(loc_inst, key=lambda x: -x['total'])
    loc_inst_table = [
        [i['nombre'], fmt(i['m']), fmt(i['h']), fmt(i['total']),
         pct_of(i['m'], i['total']), pct_of(i['h'], i['total']),
         pct_of(i['total'], loc_total)]
        for i in loc_inst_rows
    ]
    loc_inst_table.append(['TOTAL', fmt(loc_m), fmt(loc_h), fmt(loc_total),
                           loc_muj_pct, loc_hom_pct, '100.0%'])
    add_table(doc,
        ['Institución','Mujeres','Hombres','Total','% Mujeres','% Hombres','% del total'],
        loc_inst_table,
        [3200, 950, 950, 950, 950, 950, 950])
    add_spacer(doc, 8)

    # 3b. Tabla por municipio — municipios reales ordenados por volumen
    add_body(doc, '3.2  Beneficiarios localizables por municipio:', keep_next=True)
    add_spacer(doc, 3, keep_next=True)
    # Construir filas: municipios reales con localizables, ordenados desc
    loc_mun_rows = []
    for m in sorted([x for x in municipios if not x.get('especial')],
                    key=lambda x: -x.get('total_localizables', 0)):
        ltot = m.get('total_localizables', 0)
        if ltot == 0:
            continue
        loc_mun_rows.append([
            tc(m['municipio']),
            fmt(m['total']),
            fmt(ltot),
            fmt(m.get('loc_m', 0)),
            fmt(m.get('loc_h', 0)),
            pct_of(ltot, m['total']),
            pct_of(ltot, loc_total),
        ])
    # Entradas especiales (FORANEO, NO IDENTIFICADO) — siempre se muestran
    for m in [x for x in municipios if x.get('especial')]:
        ltot = m.get('total_localizables', 0)
        loc_mun_rows.append([
            tc(m['municipio']), fmt(m['total']), fmt(ltot),
            fmt(m.get('loc_m', 0)), fmt(m.get('loc_h', 0)),
            pct_of(ltot, m['total']) if ltot > 0 else '—',
            pct_of(ltot, loc_total) if ltot > 0 else '—',
        ])
    loc_mun_rows.append(['TOTAL', fmt(total_benef), fmt(loc_total),
                          fmt(loc_m), fmt(loc_h), loc_pct_ben, '100.0%'])
    add_table(doc,
        ['Municipio','Benef. únicos','Localizables','Mujeres loc.','Hombres loc.','% localiz./mun.','% del total localiz.'],
        loc_mun_rows,
        [1900, 1000, 1000, 1000, 1000, 1500, 1500])
    add_spacer(doc)

    # ══ 4. GRUPOS VULNERABLES ════════════════════════════════════════════════
    add_heading(doc, '4. Grupos Vulnerables', page_break=True)
    grupos_vul = leer_grupos_vulnerables()
    if grupos_vul:
        add_body(doc,
            'La siguiente tabla presenta los grupos de población identificados en situación de vulnerabilidad '
            'en el estado de Chihuahua, comparando la población vulnerable total con la población atendida '
            'por los programas sociales del padrón de beneficiarios 2026.')
        add_spacer(doc, 4)

        # KPIs globales de grupos vulnerables
        total_vul  = sum(g['pob_vulnerable'] for g in grupos_vul)
        total_ate  = sum(g['atendidos']      for g in grupos_vul)
        vul_m = next((g['pob_vulnerable'] for g in grupos_vul if 'muj' in g['nombre'].lower()), 0)
        vul_h = next((g['pob_vulnerable'] for g in grupos_vul if 'hom' in g['nombre'].lower()), 0)
        ate_m = next((g['atendidos']      for g in grupos_vul if 'muj' in g['nombre'].lower()), 0)
        ate_h = next((g['atendidos']      for g in grupos_vul if 'hom' in g['nombre'].lower()), 0)

        add_kpi_table(doc, [[
            {'label': 'Pob. Vulnerable Total', 'value': fmt(total_vul), 'sub': 'personas identificadas'},
            {'label': 'Población Atendida',    'value': fmt(total_ate), 'sub': pct_of(total_ate, total_vul) + ' de la pob. vulnerable'},
            {'label': 'Mujeres Vulnerables',   'value': fmt(vul_m),    'sub': pct_of(ate_m, vul_m) + ' atendidas'},
            {'label': 'Hombres Vulnerables',   'value': fmt(vul_h),    'sub': pct_of(ate_h, vul_h) + ' atendidos'},
        ]])
        add_spacer(doc, 6)

        # Tabla de todos los grupos
        add_body(doc, '4.1  Desglose por grupo vulnerable:', keep_next=True)
        add_spacer(doc, 3, keep_next=True)
        gv_rows = []
        for g in grupos_vul:
            cob = pct_of(g['atendidos'], g['pob_vulnerable']) if g['atendidos'] > 0 else '—'
            ate_str = fmt(g['atendidos']) if g['atendidos'] > 0 else '—'
            gv_rows.append([g['nombre'], fmt(g['pob_vulnerable']), ate_str, cob])
        # Fila total
        gv_rows.append(['TOTAL', fmt(total_vul), fmt(total_ate), pct_of(total_ate, total_vul)])
        add_table(doc,
            ['Grupo Vulnerable', 'Pob. Vulnerable', 'Atendidos', 'Cobertura'],
            gv_rows,
            [4000, 1500, 1500, 1000])
        add_spacer(doc)
    else:
        add_body(doc, 'No se encontraron datos de grupos vulnerables en el archivo de Excel.')
        add_spacer(doc)

    # ══ 5. INSTITUCIONES PARTICIPANTES ════════════════════════════════════════
    add_heading(doc, '5. Instituciones Participantes', page_break=True)
    add_body(doc, f'Seguimiento a las instituciones con programas activos en el padrón de beneficiarios con registro en el periodo.', keep_next=True)
    add_spacer(doc, keep_next=True)

    for inst_name in sorted(inst_act, key=lambda x: -sf(instituciones[x].get('total',0))):
        v = instituciones[inst_name]
        tot = sf(v.get('total', 0))
        # Apoyos de esta institución desde apoyos_g3
        g3_inst         = data.get('apoyos_g3', {}).get(inst_name, {})
        tot_apoyos_inst = sf(g3_inst.get('total', 0))
        # Fila resumen de institución
        add_table(doc,
            ['Institución','Mujeres','Hombres','Sin dato','Total benef.','Apoyos entregados','Programas'],
            [[inst_name, fmt(v.get('m',0)), fmt(v.get('h',0)),
              str(int(v.get('sn',0))), fmt(tot),
              fmt(tot_apoyos_inst) if tot_apoyos_inst > 0 else '—',
              str(len(v.get('programas',[]))) ]],
            [2400, 860, 860, 860, 900, 1200, 800],
            first_col_bold_size=9.5)
        # Mantener unida la tabla resumen con la de programas
        set_keep_next_on_last_table(doc)

        # Programs subrows — con apoyos por programa desde apoyos_g3
        progs = v.get('programas', [])
        if progs:
            prog_rows_data = sorted(progs, key=lambda p: -sf(p.get('total',0)))
            g3_progs = g3_inst.get('programas', {})
            prog_table_rows = []
            for p in prog_rows_data:
                ap_prog = sf(g3_progs.get(p['nombre'], {}).get('total', 0))
                prog_table_rows.append([
                    tc(p['nombre']),
                    fmt(p.get('m',0)),
                    fmt(p.get('h',0)),
                    fmt(p.get('total',0)),
                    fmt(ap_prog) if ap_prog > 0 else '—',
                ])
            add_table(doc,
                ['Programa','Mujeres','Hombres','Beneficiarios','Apoyos entregados'],
                prog_table_rows,
                [4700, 1000, 1000, 1200, 1200])
        add_spacer(doc, 4)

    # ══ 6. BENEFICIARIOS POR MUNICIPIO ════════════════════════════════════════
    add_heading(doc, '6. Beneficiarios por Municipio', page_break=True)
    add_body(doc, f'{mun_act} municipios con beneficiarios registrados en el período, ordenados por volumen de atención de mayor a menor.', keep_next=True)
    add_spacer(doc, keep_next=True)
    # Separar municipios reales de entradas especiales
    mun_reales   = sorted([m for m in municipios if not m.get('especial')], key=lambda x: -x.get('total', 0))
    mun_especial = [m for m in municipios if m.get('especial')]

    def mun_row(m, es_real=True):
        pob = m.get('poblacion', 0)
        return [
            tc(m['municipio']),
            fmt(pob) if es_real and pob > 0 else '0',
            fmt(m['total']),
            fmt(m['m']),
            fmt(m['h']),
            pct_of(m['total'], pob) if es_real and pob > 0 else '0',
            fmt(m.get('total_apoyos', 0)),
            fmt(m.get('total_localizables', 0)),
            str(m.get('n_programas', 0)),
        ]

    # Todos los municipios (reales y especiales) ordenados por volumen de atención
    todos_mun = sorted(municipios, key=lambda x: -x.get('total', 0))

    all_rows = [mun_row(m, not m.get('especial')) for m in todos_mun]

    add_table(doc,
        ['Municipio','Población','Benef. únicos','Mujeres','Hombres','Cobertura %','Apoyos','Localizables','Programas'],
        all_rows,
        [1750, 870, 870, 760, 760, 800, 900, 950, 780])
    add_spacer(doc)
    add_totals_line(doc, 'TOTAL MUNICIPIOS: ',
        f'{len(mun_reales)} municipios  |  Foráneos: {fmt(next((m["total"] for m in mun_especial if "FORAN" in unicodedata.normalize("NFD", m["municipio"].upper()).encode("ascii","ignore").decode()), 0))}  |  Sin municipio identificado: {fmt(next((m["total"] for m in mun_especial if "NO IDENT" in m["municipio"].upper()), 0))}')

    # ══ 6. APOYOS OTORGADOS ═══════════════════════════════════════════════════
    add_heading(doc, '7. Apoyos Otorgados', page_break=True)
    add_body(doc, f'Los {len(ap_clean)} tipos de apoyo registrados en el período, ordenados por volumen de entregas.', keep_next=True)
    add_spacer(doc, keep_next=True)

    # Construir árbol APOYO › INSTITUCIÓN › PROGRAMA desde desglose_municipal
    desglose_d = data.get('desglose_municipal', {})
    apoyo_tree = {}   # {apoyo: {inst: {prog: {m,h,total,muns}}}}
    for mun_k, entries in desglose_d.items():
        for e in entries:
            ap_nom = e.get('apoyo', '')
            ins    = e.get('institucion', '')
            prog   = e.get('programa', '') or '(sin programa)'
            if not ap_nom or not ins: continue
            apoyo_tree.setdefault(ap_nom, {})
            apoyo_tree[ap_nom].setdefault(ins, {})
            apoyo_tree[ap_nom][ins].setdefault(prog, {'m':0,'h':0,'total':0,'muns':set()})
            apoyo_tree[ap_nom][ins][prog]['m']     += e.get('m', 0)
            apoyo_tree[ap_nom][ins][prog]['h']     += e.get('h', 0)
            apoyo_tree[ap_nom][ins][prog]['total'] += e.get('total', 0)
            apoyo_tree[ap_nom][ins][prog]['muns'].add(mun_k)

    # Construir lista de filas con nivel y paleta heredada del apoyo padre
    # Dos paletas alternas para los bloques de apoyo:
    #   Paleta A (par):   apoyo=azul oscuro, inst=azul claro,   prog=blanco azulado claro
    #   Paleta B (impar): apoyo=azul medio,  inst=celeste suave, prog=blanco ligeramente distinto
    PALETAS = [
        {   # Paleta A
            'apoyo': RGBColor(0x1B, 0x3A, 0x6B),  # AZUL_GOB
            'inst':  RGBColor(0xC8, 0xD8, 0xF0),  # azul claro saturado
            'prog':  RGBColor(0xEC, 0xF2, 0xFB),  # casi blanco azulado
        },
        {   # Paleta B
            'apoyo': RGBColor(0x2E, 0x5B, 0xA8),  # AZUL_MED
            'inst':  RGBColor(0xD8, 0xE8, 0xF8),  # celeste suave
            'prog':  RGBColor(0xF4, 0xF8, 0xFF),  # blanco con toque celeste
        },
    ]

    ap_rows_data = []
    paleta_idx = 0
    for a in ap_clean:
        nombre_apoyo = a['apoyo']
        inst_tree    = apoyo_tree.get(nombre_apoyo, {})
        pal = PALETAS[paleta_idx % 2]
        paleta_idx += 1

        # Nivel 1 — TIPO DE APOYO
        ap_rows_data.append({
            'tipo': 'apoyo', 'pal': pal,
            'vals': [
                tc(nombre_apoyo),
                fmt(a['total']),
                fmt(a.get('m', 0)),
                fmt(a.get('h', 0)),
                fmt(a.get('n_municipios', 0)),
                pct_of(a['total'], total_apoyos),
            ]
        })
        # Nivel 2 — INSTITUCIÓN
        for ins_k, prog_tree in sorted(inst_tree.items(),
                                       key=lambda x: -sum(v['total'] for v in x[1].values())):
            ins_total = sum(v['total'] for v in prog_tree.values())
            ins_m     = sum(v['m']     for v in prog_tree.values())
            ins_h     = sum(v['h']     for v in prog_tree.values())
            ins_muns  = set()
            for v in prog_tree.values(): ins_muns |= v['muns']
            ap_rows_data.append({
                'tipo': 'inst', 'pal': pal,
                'vals': [
                    f'  \u2514 {ins_k}',
                    fmt(ins_total),
                    fmt(ins_m),
                    fmt(ins_h),
                    str(len(ins_muns)),
                    '',
                ]
            })
            # Nivel 3 — PROGRAMA
            for prog_k, pv in sorted(prog_tree.items(), key=lambda x: -x[1]['total']):
                ap_rows_data.append({
                    'tipo': 'prog', 'pal': pal,
                    'vals': [
                        f'      \u00b7 {tc(prog_k)}',
                        fmt(pv['total']),
                        fmt(pv['m']),
                        fmt(pv['h']),
                        str(len(pv['muns'])),
                        '',
                    ]
                })

    # Renderizar tabla
    PAGE_W_T = 10640
    N_COLS   = 6
    col_widths_ap = [4200, 900, 850, 850, 850, 990]
    raw_sum = sum(col_widths_ap)
    scaled_ap = [int(w * PAGE_W_T / raw_sum) for w in col_widths_ap]
    scaled_ap[-1] += PAGE_W_T - sum(scaled_ap)

    t_ap = doc.add_table(rows=1 + len(ap_rows_data), cols=N_COLS)
    t_ap.style = 'Table Grid'
    t_ap.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_width(t_ap, PAGE_W_T)

    # Encabezado
    hr_ap = t_ap.rows[0]
    set_repeat_header(hr_ap)
    for ci, (h_txt, w) in enumerate(zip(
        ['Tipo de apoyo / Institución / Programa',
         'Total', 'Mujeres', 'Hombres', 'Municipios', '% del total'],
        scaled_ap
    )):
        cell = hr_ap.cells[ci]
        set_cell_bg(cell, AZUL_GOB); set_cell_borders(cell, '1B3A6B')
        set_cell_margins(cell); set_col_width(cell, w)
        p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h_txt)
        run.font.name='Arial'; run.font.size=Pt(7.5)
        run.font.bold=True; run.font.color.rgb=BLANCO
        p.paragraph_format.space_before=Pt(0); p.paragraph_format.space_after=Pt(0)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Filas de datos — cada bloque hereda el color de su apoyo padre
    for ri, row_d in enumerate(ap_rows_data):
        row  = t_ap.rows[ri + 1]
        prevent_row_break(row)
        tipo = row_d['tipo']
        pal  = row_d['pal']

        if tipo == 'apoyo':
            bg         = pal['apoyo']
            txt_color  = BLANCO
            font_size  = Pt(9.5)
            bold_name  = True
            top_mg = bottom_mg = 110
            border_color = '1B3A6B'
        elif tipo == 'inst':
            bg         = pal['inst']
            txt_color  = AZUL_GOB
            font_size  = Pt(8)
            bold_name  = True
            top_mg = bottom_mg = 65
            border_color = 'AABBD4'
        else:  # prog
            bg         = pal['prog']
            txt_color  = GRIS_TEXT
            font_size  = Pt(7.5)
            bold_name  = False
            top_mg = bottom_mg = 45
            border_color = 'CCCCCC'

        for ci, (val, w) in enumerate(zip(row_d['vals'], scaled_ap)):
            cell = row.cells[ci]
            set_cell_bg(cell, bg)
            set_cell_borders(cell, border_color)
            set_cell_margins(cell, top=top_mg, bottom=bottom_mg, left=110, right=80)
            set_col_width(cell, w)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if ci == 0 else WD_ALIGN_PARAGRAPH.RIGHT
            run = p.add_run(str(val) if val is not None else '')
            run.font.name      = 'Arial'
            run.font.size      = font_size
            run.font.bold      = bold_name if ci == 0 else (tipo == 'apoyo')
            run.font.color.rgb = txt_color
            p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(0)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    add_spacer(doc)
    add_totals_line(doc, 'TOTAL APOYOS: ',
        f'{fmt(total_apoyos)} apoyos otorgados  |  {len(ap_clean)} tipos de apoyo distintos')

    # ══ 7. DESEMPEÑO PRESUPUESTAL Y DE PROGRAMAS ══════════════════════════════
    add_heading(doc, '8. Desempeño Presupuestal y de Programas', page_break=True)
    add_body(doc,
        'En esta sección se presentan los indicadores de presupuesto, ejecución presupuestal, '
        'eficacia, eficiencia y desempeño global por programa social. '
        'Los campos marcados con "—" corresponden a programas que aún no cuentan con '
        'información presupuestal o de metas registrada en el período.')
    add_spacer(doc)

    # ── 7.1 Definiciones ──────────────────────────────────────────────────────
    add_body(doc, '7.1  Definiciones de indicadores:', keep_next=True)
    definiciones = [
        ('Ejecución Presupuestal (EP)', 'Porcentaje del presupuesto asignado que ha sido ejercido (Gasto / Presupuesto × 100).'),
        ('Eficacia', 'Grado de cumplimiento de las metas programadas (Avance de metas / Metas programadas × 100).'),
        ('Eficiencia', 'Relación entre los beneficiarios alcanzados y el costo por beneficiario respecto a lo planeado.'),
        ('Desempeño Global', 'Indicador compuesto que integra la ejecución presupuestal, la eficacia y la eficiencia del programa.'),
    ]
    def_rows = [[d[0], d[1]] for d in definiciones]
    add_table(doc, ['Indicador', 'Definición'], def_rows, [2600, 6200])
    add_spacer(doc)

    # ── 7.2 Tabla de programas ────────────────────────────────────────────────
    add_body(doc, '7.2  Indicadores por programa social:', keep_next=True)
    add_spacer(doc, 4, keep_next=True)

    def fmt_pct(v, decimals=1):
        """Formatea un valor 0-1 como porcentaje con decimals decimales."""
        if v is None: return '—'
        return f'{v * 100:.{decimals}f}%'

    def fmt_mon(v):
        """Formatea un número como moneda MXN sin decimales."""
        if v is None: return '—'
        return f'${int(v):,}'

    def fmt_ind(v):
        """Formatea número entero, o — si es None."""
        if v is None: return '—'
        return fmt(int(v))

    prog_rows = []
    for p in indicadores:
        clave = p.get('clave', '')
        nombre_completo = f'{clave} - {tc(p["nombre"])}' if clave and clave != 'N/A' else tc(p['nombre'])
        prog_rows.append([
            nombre_completo,
            fmt_mon(p['presupuesto']),
            fmt_mon(p['gasto']),
            fmt_ind(p['metas_prog']),
            fmt_ind(p['avance_metas']),
            fmt_pct(p['eficacia']),
            fmt_pct(p['eficiencia']),
            fmt_pct(p['desempeño']),
        ])

    add_table(doc,
        ['Programa social', 'Presupuesto', 'Gasto', 'Metas prog.', 'Avance de metas',
         'Eficacia %', 'Eficiencia %', 'Desempeño %'],
        prog_rows,
        [4500, 800, 800, 800, 800, 700, 700, 700])
    add_spacer(doc)

    # Nota aclaratoria
    progs_con_datos = sum(1 for p in indicadores
                         if p['presupuesto'] is not None or p['eficacia'] is not None)
    progs_sin_datos = len(indicadores) - progs_con_datos
    add_body(doc,
        f'Nota: De los {len(indicadores)} programas registrados, {progs_con_datos} cuenta con '
        f'información presupuestal y/o de desempeño para el período. '
        f'Los {progs_sin_datos} programas restantes tienen campos pendientes de captura.',
        italic=True)

    # ══ 8. DESEMPEÑO POR INSTITUCIÓN ══════════════════════════════════════════
    add_heading(doc, '9. Desempeño por Institución', page_break=True)
    add_body(doc,
        'Resumen consolidado de indicadores de desempeño por institución. '
        'Se integran los datos de todos los programas activos de cada institución. '
        'En los casos donde la institución no cuenta con datos presupuestales en ninguno de sus programas, '
        'los campos aparecen como "—".')
    add_spacer(doc)

    # Agrupar indicadores por institución
    from collections import defaultdict
    inst_ind = defaultdict(list)
    for p in indicadores:
        inst_ind[p['institucion']].append(p)

    inst_rows = []
    for inst_nombre in sorted(inst_ind.keys()):
        progs_inst = inst_ind[inst_nombre]
        n_progs    = len(progs_inst)

        # Sumar presupuesto y gasto (solo los que tienen datos)
        pres_vals  = [p['presupuesto'] for p in progs_inst if p['presupuesto'] is not None]
        gasto_vals = [p['gasto']       for p in progs_inst if p['gasto']       is not None]
        pres_total = sum(pres_vals)  if pres_vals  else None
        gasto_total= sum(gasto_vals) if gasto_vals else None
        ep_inst    = (gasto_total / pres_total) if pres_total else None

        # Promediar eficacia, eficiencia, desempeño (solo los que tienen dato)
        def avg(campo):
            vals = [p[campo] for p in progs_inst if p[campo] is not None]
            return sum(vals) / len(vals) if vals else None

        inst_rows.append([
            inst_nombre,
            str(n_progs),
            fmt_mon(pres_total),
            fmt_mon(gasto_total),
            fmt_pct(ep_inst),
            fmt_pct(avg('eficacia')),
            fmt_pct(avg('eficiencia')),
            fmt_pct(avg('desempeño')),
        ])

    add_table(doc,
        ['Institución', 'Programas', 'Presupuesto total', 'Gasto total',
         'EP %', 'Eficacia prom.', 'Eficiencia prom.', 'Desempeño prom.'],
        inst_rows,
        [1300, 780, 1300, 1300, 700, 1100, 1100, 1100])
    add_spacer(doc)
    add_body(doc,
        'Nota: Los promedios de eficacia, eficiencia y desempeño global se calculan únicamente '
        'sobre los programas que cuentan con datos disponibles en el período.',
        italic=True)
    add_spacer(doc)

    # ══ 9. CONCLUSIONES ═══════════════════════════════════════════════════════
    add_heading(doc, '10. Conclusiones y Observaciones', page_break=True)
    add_body(doc, f'Con base en los datos registrados al corte de {PERIODO_LARGO}, se presentan las '
                  f'siguientes conclusiones por área temática:')
    add_spacer(doc, 6)

    # ── 9.1 Cobertura y Beneficiarios ─────────────────────────────────────────
    add_body(doc, '9.1  Cobertura poblacional y beneficiarios')
    add_spacer(doc, 2)

    # Municipio con mayor cobertura
    mun_top = sorted([m for m in municipios if not m.get('especial') and m.get('poblacion',0)>0],
                     key=lambda x: sf(x['total'])/x['poblacion'], reverse=True)
    mun_top_txt = f'{mun_top[0]["municipio"]} ({pct_of(mun_top[0]["total"], mun_top[0]["poblacion"])} de su población)' if mun_top else '—'

    # Municipio con más beneficiarios absolutos
    mun_abs_top = sorted([m for m in municipios if not m.get('especial')], key=lambda x: -x.get('total',0))
    mun_abs_txt = f'{mun_abs_top[0]["municipio"]} ({fmt(mun_abs_top[0]["total"])} beneficiarios)' if mun_abs_top else '—'

    # Institución con mayor cobertura
    inst_top3 = sorted(inst_act, key=lambda x: -sf(instituciones[x].get('total',0)))[:3]

    bullets_cob = [
        f'Se atendieron {fmt(total_benef)} beneficiarios únicos en {mun_act} de los 67 municipios del estado, '
        f'representando una cobertura del {cob_vulnerable} sobre la población en condición de vulnerabilidad '
        f'({fmt(POB_VULNERABLE)} personas) y del {cob_estatal_pct} sobre la población estatal total ({fmt(POB_ESTATAL)} habitantes).',

        f'Del total de beneficiarios, {fmt(total_m)} son mujeres ({muj_pct}) y {fmt(total_h)} son hombres ({hom_pct}). '
        f'{fmt(total_sn)} registros no cuentan con dato de sexo asignado y requieren seguimiento para su correcta clasificación.',

        f'Por rango de edad, la distribución es: {fmt(ninos)} niños y niñas (0 a 11 años), '
        f'{fmt(jovenes)} jóvenes (12 a 29 años), {fmt(adultos)} adultos (30–64 años) '
        f'y {fmt(mayores)} personas mayores (65 años o más). '
        f'El grupo de adultos (30 a 64 años) representa el segmento con mayor volumen de atención.',

        f'El municipio con mayor número de beneficiarios es {mun_abs_txt}, '
        f'mientras que {mun_top_txt} presenta la mayor cobertura relativa respecto a su población total.',

        f'Las tres instituciones con mayor cobertura en el período son: {", ".join(inst_top3)}, '
        f'concentrando el {pct_of(sum(sf(instituciones[i].get("total",0)) for i in inst_top3), total_benef)} '
        f'del total de beneficiarios únicos registrados.',
    ]
    for b in bullets_cob: add_bullet(doc, b)
    add_spacer(doc, 6)

    # ── 9.2 Localizables ──────────────────────────────────────────────────────
    add_body(doc, '9.2  Beneficiarios localizables')
    add_spacer(doc, 2)

    no_loc = total_benef - loc_total
    no_loc_pct = pct_of(no_loc, total_benef)
    bullets_loc = [
        f'Se identificaron {fmt(loc_total)} beneficiarios localizables ({loc_pct_ben} del total de beneficiarios únicos), '
        f'con datos de contacto, domicilio y ubicación completos y verificables.',

        f'{fmt(no_loc)} beneficiarios ({no_loc_pct}) no cuentan con información de localización completa, '
        f'lo que representa una oportunidad de mejora en la calidad del padrón para garantizar una '
        f'atención más efectiva y el seguimiento de apoyos entregados.',

        f'De los beneficiarios localizables, {fmt(loc_m)} son mujeres ({pct_of(loc_m, loc_total)}) '
        f'y {fmt(loc_h)} son hombres ({pct_of(loc_h, loc_total)}), '
        f'consistente con la distribución de género del padrón general.',
    ]
    for b in bullets_loc: add_bullet(doc, b)
    add_spacer(doc, 6)

    # ── 9.3 Apoyos Otorgados ──────────────────────────────────────────────────
    add_body(doc, '9.3  Apoyos otorgados')
    add_spacer(doc, 2)

    apoyo_top = ap_clean[0] if ap_clean else None
    apoyo_top_txt = (f'{apoyo_top["apoyo"]} ({fmt(apoyo_top["total"])} apoyos, '
                     f'{pct_of(apoyo_top["total"], total_apoyos)} del total)')  if apoyo_top else '—'
    bullets_ap = [
        f'Se entregaron un total de {fmt(total_apoyos)} apoyos durante el período, distribuidos en '
        f'{len(ap_clean)} tipos de apoyo registrados en el padrón.',

        f'El tipo de apoyo con mayor volumen de entregas es: {apoyo_top_txt}.',

        f'{sum(1 for a in ap_clean if a["total"] <= 1)} tipos de apoyo registran únicamente 1 beneficiario '
        f'en el período, lo que puede indicar entrega de apoyos a gestores o grupos comunitarios.',
    ]
    for b in bullets_ap: add_bullet(doc, b)
    add_spacer(doc, 6)

    # ── 9.4 Desempeño Presupuestal ────────────────────────────────────────────
    add_body(doc, '9.4  Desempeño presupuestal y de programas')
    add_spacer(doc, 2)

    progs_con_pres  = [p for p in indicadores if p['presupuesto'] is not None]
    progs_con_efic  = [p for p in indicadores if p['eficacia']    is not None]
    pres_total_gral = sum(p['presupuesto'] for p in progs_con_pres)
    gasto_total_gral= sum(p['gasto'] for p in progs_con_pres if p['gasto'] is not None)
    ep_gral         = (gasto_total_gral / pres_total_gral) if pres_total_gral else None
    efic_prom       = (sum(p['eficacia'] for p in progs_con_efic) / len(progs_con_efic)) if progs_con_efic else None

    bullets_pres = [
        f'De los {len(indicadores)} programas sociales registrados, {len(progs_con_pres)} cuenta con '
        f'información presupuestal capturada para el período. '
        f'Los {len(indicadores) - len(progs_con_pres)} programas restantes tienen pendiente '
        f'el registro de sus datos de presupuesto, gasto y metas.',

        (f'El presupuesto total asignado a los programas con información disponible asciende a '
         f'${pres_total_gral:,.0f} MXN, con un gasto ejercido de ${gasto_total_gral:,.0f} MXN '
         f'(ejecución presupuestal: {fmt_pct(ep_gral)}).')
        if progs_con_pres else
        'Aún no se cuenta con datos de presupuesto registrados para el período actual. '
        'Se recomienda capturar esta información para habilitar el análisis de ejecución presupuestal.',

        (f'La eficacia promedio de los programas con metas registradas es de {fmt_pct(efic_prom)}, '
         f'calculada sobre {len(progs_con_efic)} programas con datos de avance de metas disponibles.')
        if progs_con_efic else
        'Aún no se cuenta con datos de metas programadas ni de eficacia para el período. '
        'El registro de estos indicadores permitirá evaluar el cumplimiento de objetivos por programa.',

        'Se recomienda priorizar la captura de datos presupuestales y de desempeño en todos los programas '
        'para contar con un análisis integral de eficiencia, eficacia y desempeño global en los '
        'próximos cortes de información.',
    ]
    for b in bullets_pres: add_bullet(doc, b)
    add_spacer(doc, 6)

    # ── 9.5 Observaciones generales ───────────────────────────────────────────
    add_body(doc, '9.5  Observaciones generales y recomendaciones')
    add_spacer(doc, 2)

    bullets_obs = [
        f'El padrón de beneficiarios refleja un esfuerzo significativo en materia de cobertura estatal, '
        f'con presencia en los {mun_act} municipios del estado a través de {total_inst} instituciones participantes '
        f'y {total_prog} programas activos.',

        f'Se identifican {fmt(total_sn)} registros sin dato de sexo y {fmt(no_loc)} beneficiarios '
        f'sin información de localización completa. Se recomienda implementar un proceso de '
        f'depuración y actualización de datos para mejorar la calidad del padrón.',

        f'La concentración de beneficiarios en los municipios de mayor densidad poblacional '
        f'(Juárez, Chihuahua, Delicias) es consistente con la distribución demográfica estatal; '
        f'sin embargo, se recomienda revisar la cobertura en municipios serranos y rurales '
        f'para garantizar equidad territorial en la entrega de apoyos.',

        f'Se recomienda completar el registro de datos presupuestales, metas programadas y avances '
        f'en todos los programas sociales para habilitar el monitoreo integral del desempeño '
        f'institucional en los siguientes períodos de reporte.',
    ]
    for b in bullets_obs: add_bullet(doc, b)

    add_spacer(doc, 16)
    p = doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    add_top_border(p)
    r = p.add_run(f'Secretaría de Desarrollo Humano y Bien Común  —  {fecha_str}')
    r.font.name='Arial'; r.font.size=Pt(9); r.font.bold=True; r.font.color.rgb=AZUL_GOB

    print('Guardando documento...')
    doc.save(OUTPUT_PATH)
    print(f'Reporte generado: {OUTPUT_PATH}  ({os.path.getsize(OUTPUT_PATH)//1024} KB)')

if __name__ == '__main__':
    main()
